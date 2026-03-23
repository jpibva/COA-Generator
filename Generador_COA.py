import os
import re
import copy
import pdfplumber
import logging
import traceback
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import subprocess
import sys
from datetime import datetime

from coa_formats import DEFAULT_BRAND_MAP, DEFAULT_CLIENT_MAP

from coa_storage import (
    CONFIG_FILE,
    COA_REGISTRY,
    MICRO_HISTORY,
    SESSION_FILE,
    get_registro_path as _storage_get_registro_path,
    load_config,
    load_micro_history,
    load_session_data,
    normalize_product_name,
    registrar_coa as _storage_registrar_coa,
    save_config,
    save_micro_history_record as _storage_save_micro_history_record,
    save_session_data,
)

# ============================================================
# CONFIGURACIÓN
# ============================================================
APP_VERSION = "2.6"
APP_DIR = os.path.dirname(os.path.abspath(__file__))

config = load_config(CONFIG_FILE)

logging.basicConfig(
    level=logging.ERROR,
    format='%(asctime)s - %(levelname)s - %(message)s',
    filename='error_log.txt', filemode='w'
)


def _normalizar_producto(nombre):
    return normalize_product_name(nombre)


def save_micro_history_record(cliente, producto, lote, micro_values, formato_nombre=""):
    return _storage_save_micro_history_record(
        cliente,
        producto,
        lote,
        micro_values,
        config,
        formato_nombre=formato_nombre,
        history_file=MICRO_HISTORY,
    )


def get_registro_path():
    return _storage_get_registro_path(APP_DIR)


def registrar_coa(filas):
    return _storage_registrar_coa(filas, APP_DIR)

# ============================================================
# TIPOS DE TABLA MICRO
# ============================================================
MICRO_TYPE_SIMPLE    = "simple"
MICRO_TYPE_CAMERICAN = "camerican"
MICRO_TYPE_N_SERIES  = "n_series"
MICRO_TYPE_SMALL     = "small"
MICRO_TYPE_UNKNOWN   = "unknown"

def detect_micro_table_type(table):
    if not table.rows:
        return MICRO_TYPE_UNKNOWN
    first_row = [c.text.strip().lower() for c in table.rows[0].cells]
    all_text  = " ".join(c.text.strip().lower() for row in table.rows for c in row.cells)
    n_cols    = len(table.columns)
    n_rows    = len(table.rows)

    # Tablita chica: solo HepA/Norovirus, máximo 4 filas, sin TPC
    if n_rows <= 4 and ("hepatitis" in all_text or "norovirus" in all_text) and "total plate" not in all_text:
        return MICRO_TYPE_SMALL
    # Camerican: 4 col con "method" en encabezado (debe ir ANTES de n_series
    # porque Camerican también tiene "lot:" en su encabezado)
    if n_cols == 4 and any("method" in t for t in first_row):
        return MICRO_TYPE_CAMERICAN
    # n_series: tiene "lot:" en encabezado y NO tiene "method"
    if any("lot:" in t for t in first_row):
        return MICRO_TYPE_N_SERIES
    # Simple: 3+ col con microbiology o total plate
    if n_cols >= 3 and ("microbiology" in all_text or "total plate" in all_text):
        return MICRO_TYPE_SIMPLE
    return MICRO_TYPE_UNKNOWN

def find_micro_tables(doc):
    results = []
    for i, table in enumerate(doc.tables):
        all_text = " ".join(c.text.strip().lower() for row in table.rows for c in row.cells)
        if any(kw in all_text for kw in ["microbiology", "total plate", "salmonella", "listeria"]):
            tipo = detect_micro_table_type(table)
            results.append((i, table, tipo))
    return results

# ============================================================
# HELPERS DE FORMATO
# ============================================================

def _set_cell_font(cell, text, font_name, font_size, bold=False,
                   color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    # Limpiar todos los párrafos existentes en la celda
    tc = cell._tc
    for p_elem in tc.findall(qn("w:p"))[1:]:
        tc.remove(p_elem)
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

# Color azul intermedio — más apagado que celeste, más claro que azul oscuro
AZUL_MICRO = (0, 84, 166)

def set_cell_tahoma_9(cell, value):
    _set_cell_font(cell, value, "Tahoma", 9)

def set_cell_micro_simple_header(cell, lote):
    """Encabezado tabla simple: Calibri 7 centrado. RESULTS\nNROLOTE sin enter al final."""
    texto = f"RESULTS\n{lote}".rstrip()
    _set_cell_font(cell, texto, "Calibri", 7, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

def set_cell_micro_simple_result(cell, value):
    """Resultado tabla simple: Tahoma 7 centrado."""
    if value:
        _set_cell_font(cell, value, "Tahoma", 7, align=WD_ALIGN_PARAGRAPH.CENTER)

def set_cell_micro_camerican_lot(cell, value):
    """Lot: en Camerican — Calibri 12 centrado."""
    _set_cell_font(cell, value, "Calibri", 12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

def set_cell_micro_large_header(cell, value):
    """Encabezado lote n-series — Calibri 14 centrado."""
    _set_cell_font(cell, value, "Calibri", 14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

def set_cell_micro_large_result(cell, value):
    """Resultado tabla grande — Calibri 10, azul intermedio, centrado."""
    if value:
        _set_cell_font(cell, value, "Calibri", 10, color=AZUL_MICRO, align=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# MAPEO DE PARÁMETROS (con sinónimos desde config)
# ============================================================

def _build_param_map():
    """Construye el mapa sinónimo → clave desde todos los formatos en config."""
    pmap = {}
    for fmt in config["micro_formats"].values():
        for param in fmt["parametros"]:
            clave     = param["clave"]
            sinonimos = param.get("sinonimos", param["nombre"])
            for sin in sinonimos.split(","):
                pmap[sin.strip().lower()] = clave
    return pmap

def _get_param_key(text):
    pmap = _build_param_map()
    t = text.strip().lower()
    # Buscar por coincidencia exacta primero, luego parcial (más largo primero)
    for k in sorted(pmap.keys(), key=len, reverse=True):
        if k in t:
            return pmap[k]
    return None

def _get_micro_value(data, key):
    if key == "YeastMold":
        # Primero buscar si viene como campo único (Camerican)
        direct = data.get("YeastMold", "")
        if direct:
            return direct
        # Si no, combinar Yeast y Mold separados
        y = data.get("Yeast", "")
        m = data.get("Mold", "")
        if y or m:
            return f"{y} / {m}"
        return ""
    return data.get(key, "")

# ============================================================
# LLENADO TABLA SIMPLE (máx 4 lotes por tabla)
# ============================================================

def _add_col_to_table(table):
    """Agrega una columna vacía al final de la tabla."""
    try:
        for row in table.rows:
            last_tc = copy.deepcopy(row.cells[-1]._tc)
            for p in last_tc.findall(qn('w:p')):
                last_tc.remove(p)
            new_p = OxmlElement('w:p')
            new_r = OxmlElement('w:r')
            new_t = OxmlElement('w:t')
            new_t.text = ""
            new_r.append(new_t)
            new_p.append(new_r)
            last_tc.append(new_p)
            row._tr.append(last_tc)
    except Exception as e:
        logging.error(f"Error _add_col_to_table: {e}")

def _clone_table_below(doc, source_table):
    """Clona la tabla LIMPIA y la inserta debajo. Devuelve wrapper de la nueva tabla."""
    src_tbl = source_table._tbl
    sep_p   = OxmlElement('w:p')
    src_tbl.addnext(sep_p)
    new_tbl = copy.deepcopy(src_tbl)
    sep_p.addnext(new_tbl)

    from docx.table import Table as DocxTable
    wrapper = DocxTable(new_tbl, doc)
    # Limpiar TODAS las filas col 2+ (incluyendo encabezado) para evitar datos del lote anterior
    for row in wrapper.rows:
        for cell in row.cells[2:]:
            tc = cell._tc
            for p_elem in tc.findall(qn("w:p"))[1:]:
                tc.remove(p_elem)
            p = cell.paragraphs[0]
            p.clear()
    return wrapper

def fill_micro_simple(doc, source_table, lotes, micro_data_per_lote):
    """Tabla simple: máx 4 lotes por tabla. Clona hacia abajo si hay más."""
    MAX_COLS = 4
    grupos   = [lotes[i:i+MAX_COLS] for i in range(0, len(lotes), MAX_COLS)]

    tablas = [source_table]
    for _ in grupos[1:]:
        tablas.append(_clone_table_below(doc, tablas[-1]))

    for tabla, grupo in zip(tablas, grupos):
        # Encabezado primer lote
        set_cell_micro_simple_header(tabla.rows[0].cells[2], grupo[0])
        # Columnas adicionales
        for lote in grupo[1:]:
            _add_col_to_table(tabla)
            set_cell_micro_simple_header(tabla.rows[0].cells[-1], lote)
        # Llenar resultados
        for row in tabla.rows[1:]:
            key = _get_param_key(row.cells[0].text)
            if not key:
                continue
            for col_idx, lote in enumerate(grupo):
                data   = micro_data_per_lote.get(lote, {})
                val    = _get_micro_value(data, key)
                target = 2 + col_idx
                if target < len(row.cells) and val:
                    set_cell_micro_simple_result(row.cells[target], val)

def fill_micro_small(table, lotes, micro_data_per_lote):
    """
    Tablita HepA/Norovirus: fija en primera página, columnas por lote.
    Siempre "Negative" — no depende del formulario porque estos valores nunca cambian.
    Encabezado: Calibri 7 centrado. Resultados: Tahoma 7 centrado. NUNCA se duplica.
    """
    if not lotes:
        return
    # Encabezado primer lote
    set_cell_micro_simple_header(table.rows[0].cells[2], lotes[0])
    # Agregar columnas para lotes adicionales
    for lote in lotes[1:]:
        _add_col_to_table(table)
        set_cell_micro_simple_header(table.rows[0].cells[-1], lote)
    # Llenar siempre "Negative" por cada lote — valores fijos
    for row in table.rows[1:]:
        if not row.cells[0].text.strip():
            continue
        for col_idx in range(len(lotes)):
            target = 2 + col_idx
            if target < len(row.cells):
                set_cell_micro_simple_result(row.cells[target], "Negative")

# ============================================================
# CAMERICAN: una tabla por lote
# ============================================================

def fill_micro_camerican_single(table, lote, micro_data):
    """
    Tabla Camerican: 4 col → Param | Spec | Result | Method
    Fila encabezado: Lot: | Parameters | RESULT | Method used
    El número de lote va en la fila donde cells[0]="Lot:" — se limpia y reescribe.
    Los resultados van en cells[2] de cada fila de parámetro.
    """
    for row in table.rows:
        if len(row.cells) < 3:
            continue
        cell0     = row.cells[0].text.strip()
        cell0_low = cell0.lower()

        # Fila de encabezado con "Lot:" → limpiar celda completa y escribir lote
        if "lot:" in cell0_low:
            # Limpiar el texto existente (puede tener lote anterior por deepcopy)
            for para in row.cells[0].paragraphs:
                for run in para.runs:
                    run.text = ""
            set_cell_micro_camerican_lot(row.cells[0], f"Lot: {lote}")
            continue

        # Filas de parámetros → buscar por cells[0] y poner en cells[2]
        key = _get_param_key(cell0)
        if not key:
            continue
        val = _get_micro_value(micro_data, key)
        if val:
            # Limpiar celda resultado antes de escribir (evita texto duplicado)
            for para in row.cells[2].paragraphs:
                for run in para.runs:
                    run.text = ""
            set_cell_micro_large_result(row.cells[2], val)

def _find_camerican_table(doc):
    """Busca la tabla Camerican de forma flexible — 4 col con method/result/lot."""
    for i, table in enumerate(doc.tables):
        all_text = " ".join(c.text.strip().lower() for row in table.rows for c in row.cells)
        first_row_text = " ".join(c.text.strip().lower() for c in table.rows[0].cells) if table.rows else ""
        n_cols = len(table.columns)
        # Criterio: tiene "method" O ("result" y "lot") y tiene 4 columnas
        if n_cols == 4 and ("method" in all_text or ("result" in first_row_text and "lot" in first_row_text)):
            if any(kw in all_text for kw in ["total plate", "salmonella", "listeria", "coliforms"]):
                return i
    return None

def duplicate_camerican_tables(doc, lotes, micro_data_per_lote):
    """Una tabla Camerican por lote, salto de página entre ellas."""
    cam_idx = _find_camerican_table(doc)
    if cam_idx is None:
        logging.error("duplicate_camerican_tables: no se encontró tabla Camerican")
        return

    original     = doc.tables[cam_idx]
    insert_after = original._tbl
    fill_micro_camerican_single(original, lotes[0], micro_data_per_lote.get(lotes[0], {}))

    for lote in lotes[1:]:
        pb_p = OxmlElement('w:p')
        pb_r = OxmlElement('w:r')
        pb   = OxmlElement('w:br')
        pb.set(qn('w:type'), 'page')
        pb_r.append(pb)
        pb_p.append(pb_r)
        insert_after.addnext(pb_p)
        # Párrafo espaciador antes de la tabla
        spacer = OxmlElement('w:p')
        pb_p.addnext(spacer)
        new_tbl = copy.deepcopy(original._tbl)
        spacer.addnext(new_tbl)
        from docx.table import Table as DocxTable
        wrapper = DocxTable(new_tbl, doc)
        fill_micro_camerican_single(wrapper, lote, micro_data_per_lote.get(lote, {}))
        insert_after = new_tbl

# ============================================================
# N-SERIES VLM: 5 muestras por parámetro, una tabla por lote
# ============================================================

def fill_micro_n_series_single(table, lote, micro_data):
    """Llena UNA tabla n-series con datos de UN lote."""
    # Encabezado
    set_cell_micro_large_header(table.rows[0].cells[0], f"Lot: {lote}")

    for row in table.rows[1:]:
        cells      = row.cells
        if len(cells) < 3:
            continue
        param_text = cells[0].text.strip().lower()
        n_label    = cells[1].text.strip().lower()

        # Listeria y Salmonella → fila única "None detected"
        if "listeria" in param_text or "salmonella" in param_text:
            set_cell_micro_large_result(cells[1], "None detected")
            continue

        n_match = re.match(r'n(\d+)', n_label)
        if not n_match:
            continue
        n_idx = int(n_match.group(1))  # 1-based

        key = _get_param_key(param_text)
        if not key:
            continue

        # Buscar valor específico de esta n: "TPC_n1", "TPC_n2", etc.
        val_key = f"{key}_n{n_idx}"
        val     = micro_data.get(val_key, "")
        if not val:
            val = micro_data.get(key, "")
        if val:
            set_cell_micro_large_result(cells[2], val)

def duplicate_n_series_tables(doc, lotes, micro_data_per_lote):
    """
    Una tabla n-series por lote, salto de página entre ellas.
    La tablita HepA/Norovirus PERMANECE en su lugar original (primera página).
    """
    n_idx = None
    for i, table in enumerate(doc.tables):
        if detect_micro_table_type(table) == MICRO_TYPE_N_SERIES:
            n_idx = i
            break
    if n_idx is None:
        return

    original     = doc.tables[n_idx]
    insert_after = original._tbl
    fill_micro_n_series_single(original, lotes[0], micro_data_per_lote.get(lotes[0], {}))

    for lote in lotes[1:]:
        pb_p = OxmlElement('w:p')
        pb_r = OxmlElement('w:r')
        pb   = OxmlElement('w:br')
        pb.set(qn('w:type'), 'page')
        pb_r.append(pb)
        pb_p.append(pb_r)
        insert_after.addnext(pb_p)
        # Párrafo espaciador antes de la tabla
        spacer = OxmlElement('w:p')
        pb_p.addnext(spacer)
        new_tbl = copy.deepcopy(original._tbl)
        spacer.addnext(new_tbl)
        from docx.table import Table as DocxTable
        wrapper = DocxTable(new_tbl, doc)
        fill_micro_n_series_single(wrapper, lote, micro_data_per_lote.get(lote, {}))
        insert_after = new_tbl

# ============================================================
# FUNCIONES GENERALES
# ============================================================

def is_organic(product_name):
    return any(kw in product_name.lower() for kw in ['organic', 'orgánico', 'org.'])

def is_special_customer(customer_name):
    if not customer_name:
        return False
    lower = customer_name.lower()
    return any(c in lower for c in ['inabata', 'sun-in', 'sun in', 'hanmi'])

def detect_palletized(text):
    """Detecta tipo de carga desde texto del PDF. Devuelve 'pallet', 'floor' o 'slipsheet'."""
    if not text:
        return None
    t = text.strip().lower()
    if any(k in t for k in ["slip sheet", "s. sheet", "s sheet", "slipsheet"]):
        return "slipsheet"
    if any(k in t for k in ["floor", "a piso", "floored"]):
        return "floor"
    if "pallet" in t:
        return "pallet"
    return None

def fill_palletized(doc, tipo):
    """
    Estructura: Palletized | YES | [celda X] | NO | [celda X] | S. Sheet | [celda X]
    Limpia las 3 celdas de valor y pone X solo en la que corresponde.
    tipo: 'pallet' → celda tras YES, 'floor' → celda tras NO, 'slipsheet' → celda tras S.Sheet
    """
    if not tipo:
        return
    for table in doc.tables:
        for row in table.rows:
            texts = [c.text.strip().upper() for c in row.cells]
            row_joined = " ".join(texts)
            if "YES" in texts and "NO" in texts and ("S.SHEET" in row_joined or "S. SHEET" in row_joined or "SLIP" in row_joined):
                # Encontrar índices de etiquetas y sus celdas de valor (siguiente celda)
                mark_indices = {}
                for i, cell in enumerate(row.cells):
                    ct = cell.text.strip().upper()
                    if ct == "YES"  and i + 1 < len(row.cells):
                        mark_indices["pallet"]    = i + 1
                    elif ct == "NO" and i + 1 < len(row.cells):
                        mark_indices["floor"]     = i + 1
                    elif ("SHEET" in ct or "SLIP" in ct) and i + 1 < len(row.cells):
                        mark_indices["slipsheet"] = i + 1

                # Limpiar todas las celdas de valor y marcar la correcta
                for key, idx in mark_indices.items():
                    val = "X" if key == tipo else ""
                    _set_cell_font(row.cells[idx], val, "Tahoma", 9,
                                   bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                return

def replace_pesticide_text(doc, producto):
    """
    FIX: Solo reemplaza párrafos que tengan la leyenda completa de pesticidas
    (que contengan 'lmr' o 'mrl' o 'chemical residues').
    NO toca el título 'PESTICIDES ANALYSIS'.
    """
    leyenda = (
        "Product Free From Chemical Residues" if is_organic(producto)
        else "Chemical Pesticides according to LMR's destination market"
    )
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    for p in all_paragraphs:
        p_lower = p.text.lower()
        # Solo reemplazar si tiene la leyenda específica (lmr, mrl o chemical residues)
        # Excluye el título "PESTICIDES ANALYSIS" que solo tiene "pesticid"
        if ("lmr" in p_lower or "mrl" in p_lower or
                ("chemical" in p_lower and "residue" in p_lower)):
            p.clear()
            run = p.add_run(leyenda)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            p.alignment   = WD_ALIGN_PARAGRAPH.CENTER
            return

def register_arrow_nav(widget_list):
    """
    Registra navegación con flechas arriba/abajo/izq/der para una lista de Entry widgets.
    - Arriba/Abajo: saltan al entry anterior/siguiente
    - Izquierda al borde izquierdo: salta al entry anterior
    - Derecha al borde derecho: salta al entry siguiente
    - Enter/Tab: avanza al siguiente
    """
    def _go(idx, delta):
        target = (idx + delta) % len(widget_list)
        w = widget_list[target]
        w.focus_set()
        try:
            w.icursor(tk.END if delta > 0 else tk.END)
        except Exception:
            pass

    for i, w in enumerate(widget_list):
        def _make(idx):
            def _key(e):
                keysym = e.keysym
                if keysym == "Up":
                    _go(idx, -1); return "break"
                if keysym == "Down" or keysym == "Return":
                    _go(idx, 1); return "break"
                if keysym == "Left":
                    try:
                        cur = widget_list[idx].index(tk.INSERT)
                        if cur == 0:
                            _go(idx, -1); return "break"
                    except Exception:
                        pass
                if keysym == "Right":
                    try:
                        w2 = widget_list[idx]
                        cur = w2.index(tk.INSERT)
                        end = len(w2.get())
                        if cur >= end:
                            _go(idx, 1); return "break"
                    except Exception:
                        pass
            return _key
        w.bind("<Key>", _make(i))

def is_camerican(cliente):
    """Detecta si el cliente es Camerican."""
    return "camerican" in (cliente or "").lower()

def read_camerican_brix_labels(doc):
    """
    Lee las filas 'Average Brix [fruta]' y 'Average pH' del template Camerican.
    Retorna lista de (label_original, key) donde key es el texto completo de la celda.
    """
    labels = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                low = txt.lower()
                if "average brix" in low or "average ph" in low or "avarage brix" in low:
                    if txt not in [l[0] for l in labels]:
                        labels.append((txt, txt))
    return labels

def _parse_limit_value(txt):
    """Extrae valor numérico de un texto de límite tipo 5.00 %, 10.00%, 2 un, etc."""
    import re as _re2
    txt = txt.strip()
    m = _re2.search(r"(\d+\.?\d*)\s*%", txt)
    if m:
        return float(m.group(1)), "%"
    m = _re2.search(r"(\d+\.?\d*)", txt)
    if m:
        return float(m.group(1)), "num"
    return None, None

def read_camerican_defect_structure(doc):
    """
    Lee la tabla DEFECTS LIMIT del template Camerican.
    Retorna:
      - frutas: lista de nombres de frutas (columnas de resultados)
      - params: lista de nombres de parámetros (filas)
      - result_col_start: índice de la primera columna de resultados
      - defaults: dict {param: {fruta: valor_default}} leído de las celdas de resultado
      - limits: dict {param: {fruta: (valor_num, tipo)}} leído de columnas de límite
    """
    for table in doc.tables:
        all_text = " ".join(c.text.strip().lower()
                            for row in table.rows for c in row.cells)
        if "defect" not in all_text:
            continue
        if len(table.rows) < 2:
            continue
        header_row = table.rows[0]
        ncols = len(header_row.cells)

        # Detectar columnas de frutas izq (límites) y der (resultados)
        seen = {}
        fruit_cols_left  = []   # [(col_idx, fruta)]
        fruit_cols_right = []   # [(col_idx, fruta)]
        for ci, cell in enumerate(header_row.cells):
            txt = cell.text.strip()
            if not txt or txt.lower() in ("defects limit", "defect limit", ""):
                continue
            if txt in seen:
                fruit_cols_right.append((ci, txt))
            else:
                seen[txt] = ci
                fruit_cols_left.append((ci, txt))

        frutas = [t for _, t in fruit_cols_left]
        result_col_map = {t: ci for ci, t in fruit_cols_right}  # fruta → col resultado
        limit_col_map  = {t: ci for ci, t in fruit_cols_left}   # fruta → col límite
        result_col_start = fruit_cols_right[0][0] if fruit_cols_right else (ncols // 2 + 1)

        params   = []
        defaults = {}  # {param: {fruta: default_val}}
        limits   = {}  # {param: {fruta: (num, tipo)}}

        for row in table.rows[1:]:
            if not row.cells:
                continue
            p = row.cells[0].text.strip()
            if not p or p.lower() in ("parameters", "% parameters", ""):
                continue
            params.append(p)
            defaults[p] = {}
            limits[p]   = {}
            for fruta in frutas:
                # Default: valor en columna de resultado
                rc = result_col_map.get(fruta)
                if rc is not None and rc < len(row.cells):
                    defaults[p][fruta] = row.cells[rc].text.strip()
                else:
                    defaults[p][fruta] = ""
                # Límite: valor en columna de límite
                lc = limit_col_map.get(fruta)
                if lc is not None and lc < len(row.cells):
                    lim_txt = row.cells[lc].text.strip()
                    num, tipo = _parse_limit_value(lim_txt)
                    limits[p][fruta] = (num, tipo, lim_txt)
                else:
                    limits[p][fruta] = (None, None, "")

        return frutas, params, result_col_start, defaults, limits
    return [], [], 1, {}, {}

def fill_camerican_brix_ph(doc, brix_ph_vals):
    """
    Llena las celdas de Brix/pH en el template Camerican.
    brix_ph_vals: dict {label_original: valor}
    """
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells[:-1]):
                txt = cell.text.strip()
                if txt in brix_ph_vals:
                    val = brix_ph_vals[txt]
                    set_cell_tahoma_9(row.cells[i + 1], val)

def fill_camerican_defects(doc, frutas, quality_matrix):
    """
    Llena la tabla DEFECTS LIMIT del template Camerican.
    quality_matrix: dict {param: {fruta: valor}}
    Solo escribe en las columnas de resultados (derecha).
    """
    for table in doc.tables:
        all_text = " ".join(c.text.strip().lower()
                            for row in table.rows for c in row.cells)
        if "defect" not in all_text:
            continue
        if len(table.rows) < 2:
            continue

        header_row = table.rows[0]
        ncols = len(header_row.cells)

        # Detectar columnas de resultado (repetidas)
        seen = {}
        result_cols = {}  # fruta → col index
        for ci, cell in enumerate(header_row.cells):
            txt = cell.text.strip()
            if not txt or txt.lower() in ("defects limit", "defect limit", ""):
                continue
            if txt in seen:
                result_cols[txt] = ci
            else:
                seen[txt] = ci

        if not result_cols:
            # fallback: mitad derecha
            mid = ncols // 2
            for j, fruta in enumerate(frutas):
                if mid + j < ncols:
                    result_cols[fruta] = mid + j

        # Llenar filas de parámetros
        for row in table.rows[1:]:
            if not row.cells:
                continue
            param = row.cells[0].text.strip()
            if not param or param.lower() in ("parameters", "% parameters"):
                continue
            param_clean = re.sub(r"[^a-zA-Z]", "", param).lower()
            for fruta, col_idx in result_cols.items():
                # buscar valor en quality_matrix por fruta y parámetro
                fruta_clean = re.sub(r"[^a-zA-Z]", "", fruta).lower()
                val = ""
                for p_key, f_dict in quality_matrix.items():
                    p_clean = re.sub(r"[^a-zA-Z]", "", p_key).lower()
                    if p_clean == param_clean:
                        for f_key, v in f_dict.items():
                            f_key_clean = re.sub(r"[^a-zA-Z]", "", f_key).lower()
                            if f_key_clean == fruta_clean:
                                val = v
                                break
                        break
                if col_idx < len(row.cells):
                    _set_cell_font(row.cells[col_idx], val, "Tahoma", 7,
                                   align=WD_ALIGN_PARAGRAPH.CENTER)
        return

def fill_defects_table(doc, formato_nombre, quality_values):
    """
    Llena la tabla DEFECTS LIMIT con los valores ingresados.
    Soporta 2 columnas (Parameters|Results) y 3 columnas (Parameters|Limit|Results).
    Fuente: Tahoma 7 centrado.
    """
    fmt    = config["quality_formats"].get(formato_nombre, {})
    params = fmt.get("parametros", [])
    ncols  = fmt.get("columnas", 2)
    if not params:
        return

    # Índice de columna de results: col 1 si 2 cols, col 2 si 3 cols
    result_col = ncols - 1

    # Mapa nombre_limpio → valor
    val_map = {re.sub(r"[^a-zA-Z]", "", p).lower(): quality_values.get(p, "")
               for p in quality_values}

    for table in doc.tables:
        all_text = " ".join(c.text.strip().lower()
                            for row in table.rows for c in row.cells)
        if "defect" not in all_text:
            continue
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = row.cells[0].text.strip()
            label_clean = re.sub(r"[^a-zA-Z]", "", label).lower()
            if not label_clean or label_clean in ("parameters", "defectslimit"):
                continue
            val = val_map.get(label_clean, "")
            if result_col < len(row.cells):
                _set_cell_font(row.cells[result_col], val, "Tahoma", 7,
                               align=WD_ALIGN_PARAGRAPH.CENTER)
        return  # solo llenar la primera tabla DEFECTS encontrada

def fill_all_tables(doc, replacements):
    """
    Busca etiquetas en CUALQUIER celda de la fila (no solo col 0).
    El valor se escribe en la celda inmediatamente siguiente.
    """
    clean = {re.sub(r"[^a-zA-Z]", "", k).lower(): v for k, v in replacements.items()}
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells[:-1]):
                key = re.sub(r"[^a-zA-Z]", "", cell.text.strip()).lower()
                if key and key in clean:
                    set_cell_tahoma_9(row.cells[i + 1], clean[key])

def is_livemore(cliente, producto_formato=""):
    """Detecta si es un COA de Livemore."""
    if not cliente:
        return False
    return "livemore" in cliente.lower() or "livemore" in producto_formato.lower()

def _calcular_exp(fecha_prod_str, años):
    """Calcula fecha de expiración dado fecha de producción y años a sumar."""
    try:
        for fmt in ("%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
            try:
                dt = datetime.strptime(fecha_prod_str, fmt)
                return dt.replace(year=dt.year + años).strftime("%d-%m-%Y")
            except ValueError:
                continue
    except Exception:
        pass
    return ""

def fill_lot_table(doc, lot_data):
    """
    Llena la tabla de identificación de lotes.
    lot_data keys: Lotes, Cantidades, Fechas, Cliente, MicroFormato
    
    Manufacture Date: lote / fecha = X cases  (todos)
    Packed Date:      lote / fecha            (todos)
    Lot Number:       lote1, lote2, ...       (todos)
    Quantity:         lote = X cases          (todos)
    Expiration Date:  lote / fecha+2años      (Livemore)
                      lote / fecha+3años      (asiáticos: Inabata, Sun-in, Hanmi)
                      vacío                   (resto)
    """
    lotes     = lot_data.get('Lotes', [])
    cantidades= lot_data.get('Cantidades', [])
    fechas    = lot_data.get('Fechas', [])
    cliente   = lot_data.get('Cliente', '')
    fmt_micro = lot_data.get('MicroFormato', '')

    if not lotes:
        return

    # Calcular líneas por campo
    manuf_lines = []
    packed_lines = []
    exp_lines    = []
    fechas_fmt   = []  # solo fechas formateadas, sin lotes ni cajas
    lot_line     = ", ".join(lotes)
    qty_lines    = []

    livemore  = is_livemore(cliente, fmt_micro)
    especial  = is_special_customer(cliente)

    for i, lote in enumerate(lotes):
        fecha  = fechas[i] if i < len(fechas) else ""
        cant   = cantidades[i] if i < len(cantidades) else "0"
        try:    cajas = int(float(cant))
        except: cajas = 0

        # Formato fecha para mostrar: DD-MM-YYYY
        fecha_fmt = fecha.replace("/", "-") if fecha else ""

        manuf_lines.append(f"{lote} / {fecha_fmt} = {cajas} cases")
        packed_lines.append(f"{lote} / {fecha_fmt}")
        fechas_fmt.append(fecha_fmt)
        qty_lines.append(f"{lote} = {cajas} cases")

        if livemore:
            exp_date = _calcular_exp(fecha, 2)
            exp_lines.append(f"{lote} / {exp_date}")

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = row.cells[0].text.strip().lower()

            if "lot number" in label or label == "lot no" or label == "lot":
                set_cell_tahoma_9(row.cells[1], lot_line)

            elif "manufacture" in label and "name" not in label:
                # Livemore: lote / fecha = X cases | Resto: solo fechas separadas por coma
                if livemore:
                    set_cell_tahoma_9(row.cells[1], "\n".join(manuf_lines))
                else:
                    set_cell_tahoma_9(row.cells[1], ", ".join(fechas_fmt))

            elif "packed" in label and "name" not in label:
                set_cell_tahoma_9(row.cells[1], "\n".join(packed_lines))
                # Buscar "Expiration date" en la misma fila (col 2) — solo Livemore
                if livemore and len(row.cells) > 3:
                    label2 = row.cells[2].text.strip().lower()
                    if ("expir" in label2 or "best before" in label2) and exp_lines:
                        set_cell_tahoma_9(row.cells[3], "\n".join(exp_lines))

            elif "production" in label and "name" not in label:
                set_cell_tahoma_9(row.cells[1], "\n".join(packed_lines))

            elif "quantity per date" in label or "qty per" in label:
                set_cell_tahoma_9(row.cells[1], "\n".join(qty_lines))

            elif ("expir" in label or "best before" in label) and exp_lines:
                set_cell_tahoma_9(row.cells[1], "\n".join(exp_lines))

def estilizar_nombre_archivo(embarque, po, producto):
    producto_fmt = re.sub(r"[^a-zA-Z0-9]+", "_", producto.strip())[:40]
    po_clean     = re.sub(r'[^0-9]', '', str(po))
    return f"COA_{embarque}_PO{po_clean}_{producto_fmt}.docx"

def format_shipment_number(raw):
    if not raw: return "SIN_EMBARQUE"
    numeric = re.sub(r'\D', '', raw)
    return f"E-{numeric.zfill(5)}" if numeric else "SIN_EMBARQUE"

def extract_data_from_pdf(pdf_path):
    data = {"general_info": {}, "products": []}
    with pdfplumber.open(pdf_path) as pdf:
        full_text = ""
        for page in pdf.pages:
            full_text += page.extract_text(x_tolerance=2, y_tolerance=3) + "\n"
        patterns = {
            "raw_embarque": r"(?:Boarding|Shipment|Embarque)\s*Nº?\s*([A-ZΕΝ\s-]*\d+)",
            "cliente":      r"Customer:\s*(.*?)(?:\n|Container)",
            "po":           r"PO\s*Number:\s*(\S+)",
            "contenedor":   r"Container\s*Number:\s*([^\n]+)",
            "sello":        r"Seal\s*Number:\s*([^\n]+)",

            "load_type":    r"(?:Load|Carga|Palletized|Tipo de carga)[:\s]+([^\n]+)",
        }
        for key, pattern in patterns.items():
            match = re.search(pattern, full_text, re.IGNORECASE)
            if match:
                data["general_info"][key] = match.group(1).strip()
        data["general_info"]["embarque"] = format_shipment_number(
            data["general_info"].get("raw_embarque", ""))

        # Detectar TODOS los códigos TZ (Ryan Recorder)
        tz_codes = re.findall(r"TZ\d+", full_text)
        if tz_codes:
            data["general_info"]["recorder"] = " - ".join(tz_codes)
        product_blocks = re.split(r'Product:', full_text, flags=re.IGNORECASE)
        for block in product_blocks[1:]:
            m = re.search(r'^\s*(.*?)\n', block)
            product_name = m.group(1).strip() if m else "Producto Desconocido"
            uw = re.search(r'Net\s*Weight\s*\(unit\).*?\n\s*([\d,.]+)\s*Kgs', block, re.IGNORECASE | re.DOTALL)
            unit_weight = uw.group(1).replace(',', '.') if uw else "0"
            lot_pat = re.compile(r"^(\d{5}(?:-\d+)?)\s+([\d,.]+)\s+([\d,.]+)\s+([\d,.]+)\s+([\d/]+)", re.MULTILINE)
            lots = []
            for match in lot_pat.finditer(block):
                lots.append({
                    "producto":         product_name,
                    "lote":             match.group(1),
                    "cantidad":         match.group(2).replace(".", "").replace(",", ""),
                    "peso_neto_total":  match.group(3).replace(".", "").replace(",", "."),
                    "peso_bruto_total": match.group(4).replace(".", "").replace(",", "."),
                    "fecha_produccion": match.group(5),
                    "peso_x_unidad":    unit_weight
                })
            if lots:
                data["products"].extend(lots)
    return data

# ============================================================
# VENTANA DE MICROBIOLOGÍA
# ============================================================

class MicrobiologyInputDialog(tk.Toplevel):
    def __init__(self, parent, productos_lotes, productos_formatos, cliente, micro_history):
        super().__init__(parent)
        self.title("Resultados Microbiológicos")
        self.geometry("820x660")
        self.resizable(True, True)
        self.grab_set()

        self.cliente            = cliente
        self.productos_lotes    = productos_lotes
        self.productos_formatos = productos_formatos
        self.micro_history      = micro_history
        self.confirmed          = False
        self.vars               = {}
        self.formato_vars       = {}
        self.lote_nb_frames     = {}
        self._clipboard         = {}  # {clave: valor} copiado de un lote

        self._build_ui()
        self.protocol("WM_DELETE_WINDOW", self._on_cancel)

    def _build_ui(self):
        # Botones fijos abajo
        btn = ttk.Frame(self)
        btn.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=8)
        ttk.Button(btn, text="✔  Confirmar y Generar", command=self._on_confirm).pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn, text="✘  Cancelar",            command=self._on_cancel).pack(side=tk.RIGHT)

        top = ttk.Frame(self)
        top.pack(fill=tk.BOTH, expand=True, padx=10, pady=(10,0))

        ttk.Label(top,
            text="Rojo = obligatorio  ·  Gris = precargado editable  ·  Vacío = sin resultado aún",
            foreground="#2A4A6B", font=("Segoe UI", 8, "italic")
        ).pack(anchor=tk.W, pady=(0, 6))

        nb = ttk.Notebook(top)
        nb.pack(fill=tk.BOTH, expand=True)
        for producto, lotes in self.productos_lotes.items():
            tab = ttk.Frame(nb)
            nb.add(tab, text=producto[:30])
            self._build_product_tab(tab, producto, lotes)

        # Rueda del mouse — apunta al canvas bajo el cursor
        def _route_mousewheel(e):
            w = e.widget
            while w:
                if isinstance(w, tk.Canvas):
                    w.yview_scroll(int(-1*(e.delta/120)), "units")
                    return
                try:
                    w = w.master
                except Exception:
                    break
        self.bind_all("<MouseWheel>", _route_mousewheel)
        self.after(150, self._apply_arrow_nav)

    def _build_product_tab(self, parent, producto, lotes):
        auto_micro = _detect_micro_format(self.cliente, producto)
        formato_actual = self.productos_formatos.get(producto, auto_micro)
        formato_var    = tk.StringVar(value=formato_actual)
        self.formato_vars[producto] = formato_var

        sel = ttk.LabelFrame(parent, text="Formato microbiológico", padding=5)
        sel.pack(fill=tk.X, padx=5, pady=5)
        combo = ttk.Combobox(sel, textvariable=formato_var,
                             values=list(config["micro_formats"].keys()),
                             state="readonly", width=30)
        combo.pack(side=tk.LEFT, padx=5)
        ttk.Button(sel, text="Aplicar",
                   command=lambda p=producto, l=lotes: self._rebuild_lote_tabs(p, l)
                   ).pack(side=tk.LEFT, padx=5)


        lote_container = ttk.Frame(parent)
        lote_container.pack(fill=tk.BOTH, expand=True)
        self.lote_nb_frames[producto] = lote_container
        self._rebuild_lote_tabs(producto, lotes)

    def _rebuild_lote_tabs(self, producto, lotes):
        container = self.lote_nb_frames[producto]
        for w in container.winfo_children():
            w.destroy()

        fmt_nombre = self.formato_vars[producto].get()
        fmt        = config["micro_formats"].get(fmt_nombre, {})
        params     = fmt.get("parametros", [])
        tipo       = fmt.get("tipo", "simple")

        nb_lotes = ttk.Notebook(container)
        nb_lotes.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.vars[producto] = {}
        for lote in lotes:
            tab = ttk.Frame(nb_lotes)
            nb_lotes.add(tab, text=f"Lote {lote}")
            self._build_lot_tab(tab, producto, lote, params, tipo)

    def _copy_lote(self, producto, lote):
        """Copia los valores del lote actual al clipboard."""
        self._clipboard = {k: v.get() for k, v in self.vars.get(producto, {}).get(lote, {}).items()}
        self._clipboard_source = f"Lote {lote}"

    def _paste_lote(self, producto, lote):
        """Pega los valores del clipboard en el lote actual."""
        if not self._clipboard:
            messagebox.showinfo("Pegar", "No hay valores copiados aún.")
            return
        lote_vars = self.vars.get(producto, {}).get(lote, {})
        for k, v in self._clipboard.items():
            if k in lote_vars:
                lote_vars[k].set(v)

    def _build_lot_tab(self, parent, producto, lote, params, tipo):
        hist_key = (self.cliente.lower().strip(), _normalizar_producto(producto), lote)
        hist     = self.micro_history.get(hist_key, {})

        # Barra superior con historial y botones copiar/pegar
        top_bar = ttk.Frame(parent)
        top_bar.pack(fill=tk.X, padx=10, pady=(5,0))

        if hist:
            ttk.Label(top_bar,
                text=f"✔ Precargado desde historial (último uso: {hist.get('Fecha', '?')})",
                foreground="green", font=('Helvetica', 9)
            ).pack(side=tk.LEFT)

        ttk.Button(top_bar, text="📋 Copiar resultados",
                   command=lambda p=producto, l=lote: self._copy_lote(p, l)
                   ).pack(side=tk.RIGHT, padx=(4,0))
        ttk.Button(top_bar, text="📌 Pegar resultados",
                   command=lambda p=producto, l=lote: self._paste_lote(p, l)
                   ).pack(side=tk.RIGHT, padx=(4,0))

        container = ttk.Frame(parent)
        container.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        sb     = ttk.Scrollbar(container, orient="vertical")
        canvas = tk.Canvas(container, highlightthickness=0, yscrollcommand=sb.set)
        sb.config(command=canvas.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        sf = ttk.Frame(canvas)
        win = canvas.create_window((0, 0), window=sf, anchor="nw")

        def _on_sf_cfg(e, c=canvas):
            c.configure(scrollregion=c.bbox("all"))
        def _on_canvas_cfg(e, c=canvas, w=win):
            c.itemconfig(w, width=e.width)
        sf.bind("<Configure>", _on_sf_cfg)
        canvas.bind("<Configure>", _on_canvas_cfg)

        self.vars[producto][lote] = {}
        row_idx = 0

        for param in params:
            nombre   = param["nombre"]
            clave    = param["clave"]
            defecto  = param.get("defecto", "")
            ingresar = param.get("ingresar", False)
            n_series = param.get("n_series", False)

            if tipo == "n_series" and n_series:
                ttk.Label(sf, text=nombre, font=('Helvetica', 9, 'bold'),
                          anchor=tk.W).grid(row=row_idx, column=0, columnspan=10,
                                            padx=5, pady=(8, 2), sticky=tk.W)
                row_idx += 1
                for n in range(1, 6):
                    val_key  = f"{clave}_n{n}"
                    hist_val = hist.get(val_key, defecto)
                    ttk.Label(sf, text=f"n{n}:", width=4, anchor=tk.E).grid(
                        row=row_idx, column=(n-1)*2, padx=(4,1), pady=1, sticky=tk.E)
                    var = tk.StringVar(value=str(hist_val) if hist_val else "")
                    entry = ttk.Entry(sf, textvariable=var, width=10)
                    entry.grid(row=row_idx, column=(n-1)*2+1, padx=(0,3), pady=1)
                    self.vars[producto][lote][val_key] = var
                row_idx += 1
            else:
                hist_val = hist.get(clave, defecto)
                ttk.Label(sf, text=nombre, width=26, anchor=tk.W).grid(
                    row=row_idx, column=0, padx=5, pady=3, sticky=tk.W)
                var = tk.StringVar(value=str(hist_val) if hist_val else "")
                entry = ttk.Entry(sf, textvariable=var, width=22)
                entry.grid(row=row_idx, column=1, padx=5, pady=3)
                self.vars[producto][lote][clave] = var
                row_idx += 1

    def _apply_format_to_all(self, formato):
        """Aplica el mismo formato micro a todos los productos y reconstruye tabs."""
        for prod, fv in self.formato_vars.items():
            fv.set(formato)
        for prod, lotes in self._productos_lotes.items():
            self._rebuild_lote_tabs(prod, lotes)
        messagebox.showinfo("Formato aplicado",
            f"Formato '{formato}' aplicado a todos los productos.")

    def _apply_arrow_nav(self):
        """Recolecta todos los Entry de la ventana y registra navegación."""
        entries = []
        def _collect(w):
            if isinstance(w, (ttk.Entry, tk.Entry)):
                entries.append(w)
            for child in w.winfo_children():
                _collect(child)
        _collect(self)
        register_arrow_nav(entries)

    def _on_confirm(self):
        self.confirmed = True
        self.destroy()

    def _on_cancel(self):
        self.confirmed = False
        self.destroy()

    def get_results(self):
        results = {}
        for producto, lotes_dict in self.vars.items():
            results[producto] = {}
            for lote, campos_dict in lotes_dict.items():
                results[producto][lote] = {k: v.get() for k, v in campos_dict.items()}
        return results

    def get_formatos(self):
        return {p: v.get() for p, v in self.formato_vars.items()}

# ============================================================
# EDITOR DE FORMATOS MICROBIOLÓGICOS
# ============================================================

class MicroFormatEditor(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("Editor de Formatos Microbiológicos")
        self.geometry("960x600")
        self.resizable(True, True)
        self.grab_set()
        self._build_ui()

    def _build_ui(self):
        main = ttk.PanedWindow(self, orient=tk.HORIZONTAL)
        main.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        left = ttk.Frame(main, width=200)
        main.add(left, weight=1)
        ttk.Label(left, text="Formatos:", font=('Helvetica', 10, 'bold')).pack(anchor=tk.W)
        lb_frame = ttk.Frame(left)
        lb_frame.pack(fill=tk.BOTH, expand=True)
        self.listbox = tk.Listbox(lb_frame, width=22, height=18, selectmode=tk.SINGLE)
        self.listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        lb_scroll = ttk.Scrollbar(lb_frame, orient="vertical", command=self.listbox.yview)
        lb_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.listbox.config(yscrollcommand=lb_scroll.set)
        self.listbox.bind("<<ListboxSelect>>", self._on_select)
        bf = ttk.Frame(left)
        bf.pack(fill=tk.X, pady=5)
        ttk.Button(bf, text="+ Nuevo",  command=self._new,   width=7).pack(side=tk.LEFT, padx=1)
        ttk.Button(bf, text="Borrar",   command=self._delete,width=7).pack(side=tk.LEFT, padx=1)
        ttk.Button(bf, text="Subir",    command=self._move_up,  width=6).pack(side=tk.LEFT, padx=1)
        ttk.Button(bf, text="Bajar",    command=self._move_down,width=6).pack(side=tk.LEFT, padx=1)

        right = ttk.Frame(main)
        main.add(right, weight=4)

        top = ttk.Frame(right)
        top.pack(fill=tk.X, pady=5)
        ttk.Label(top, text="Nombre:").grid(row=0, column=0, sticky=tk.W, padx=5)
        self.nombre_var = tk.StringVar()
        ttk.Entry(top, textvariable=self.nombre_var, width=22).grid(row=0, column=1, padx=5)
        ttk.Label(top, text="Tipo:").grid(row=0, column=2, sticky=tk.W, padx=5)
        self.tipo_var = tk.StringVar(value="simple")
        ttk.Combobox(top, textvariable=self.tipo_var,
                     values=["simple", "camerican", "n_series"],
                     state="readonly", width=12).grid(row=0, column=3, padx=5)
        ttk.Label(top, text="Descripción:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        self.desc_var = tk.StringVar()
        ttk.Entry(top, textvariable=self.desc_var, width=55).grid(row=1, column=1, columnspan=3, padx=5)

        # Encabezados columnas
        hdr = ttk.Frame(right)
        hdr.pack(fill=tk.X, padx=5)
        for i, (h, w) in enumerate([("Nombre parámetro",16),("Clave",10),
                                      ("Defecto",10),("Ingresar?",8),
                                      ("N-series?",8),("Sinónimos (separados por coma)",28)]):
            ttk.Label(hdr, text=h, font=('Helvetica', 8, 'bold'), width=w).grid(row=0, column=i, padx=2)

        # Canvas scroll para parámetros
        cf = ttk.Frame(right)
        cf.pack(fill=tk.BOTH, expand=True, padx=5)
        self.pcanvas = tk.Canvas(cf, highlightthickness=0)
        sb = ttk.Scrollbar(cf, orient="vertical", command=self.pcanvas.yview)
        self.psf = ttk.Frame(self.pcanvas)
        self.psf.bind("<Configure>", lambda e: self.pcanvas.configure(scrollregion=self.pcanvas.bbox("all")))
        self.pcanvas.create_window((0, 0), window=self.psf, anchor="nw")
        self.pcanvas.configure(yscrollcommand=sb.set)
        self.pcanvas.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        self.param_rows = []

        pb = ttk.Frame(right)
        pb.pack(fill=tk.X, padx=5, pady=3)
        ttk.Button(pb, text="+ Agregar parámetro", command=self._add_row).pack(side=tk.LEFT)

        sf2 = ttk.Frame(right)
        sf2.pack(fill=tk.X, padx=5, pady=8)
        ttk.Button(sf2, text="💾 Guardar formato", command=self._save).pack(side=tk.RIGHT)
        self.slbl = ttk.Label(sf2, text="", foreground="green")
        self.slbl.pack(side=tk.LEFT)

        self._refresh_list()

    def _refresh_list(self):
        self.listbox.delete(0, tk.END)
        for n in config["micro_formats"]:
            self.listbox.insert(tk.END, n)

    def _on_select(self, event):
        sel = self.listbox.curselection()
        if not sel: return
        name = self.listbox.get(sel[0])
        fmt  = config["micro_formats"].get(name, {})
        self.nombre_var.set(name)
        self.tipo_var.set(fmt.get("tipo", "simple"))
        self.desc_var.set(fmt.get("descripcion", ""))
        for w in self.psf.winfo_children(): w.destroy()
        self.param_rows = []
        for p in fmt.get("parametros", []):
            self._add_row(p)

    def _add_row(self, param=None):
        i  = len(self.param_rows)
        rv = {
            "nombre":   tk.StringVar(value=param.get("nombre",  "") if param else ""),
            "clave":    tk.StringVar(value=param.get("clave",   "") if param else ""),
            "defecto":  tk.StringVar(value=param.get("defecto", "") if param else ""),
            "ingresar": tk.BooleanVar(value=param.get("ingresar", False) if param else False),
            "n_series": tk.BooleanVar(value=param.get("n_series", False) if param else False),
            "sinonimos":tk.StringVar(value=param.get("sinonimos","") if param else ""),
        }
        ttk.Entry(self.psf, textvariable=rv["nombre"],   width=17).grid(row=i, column=0, padx=2, pady=2)
        ttk.Entry(self.psf, textvariable=rv["clave"],    width=10).grid(row=i, column=1, padx=2, pady=2)
        ttk.Entry(self.psf, textvariable=rv["defecto"],  width=10).grid(row=i, column=2, padx=2, pady=2)
        ttk.Checkbutton(self.psf, variable=rv["ingresar"]).grid(row=i, column=3, padx=2, pady=2)
        ttk.Checkbutton(self.psf, variable=rv["n_series"]).grid(row=i, column=4, padx=2, pady=2)
        ttk.Entry(self.psf, textvariable=rv["sinonimos"], width=28).grid(row=i, column=5, padx=2, pady=2)
        ttk.Button(self.psf, text="✕", width=2,
                   command=lambda r=rv: self._remove_row(r)).grid(row=i, column=6, padx=2)
        self.param_rows.append(rv)

    def _remove_row(self, rv):
        self.param_rows.remove(rv)
        for w in self.psf.winfo_children(): w.destroy()
        temp = list(self.param_rows)
        self.param_rows = []
        for p in temp:
            self._add_row({k: v.get() for k, v in p.items()})

    def _save(self):
        name = self.nombre_var.get().strip()
        if not name:
            messagebox.showwarning("Error", "El nombre no puede estar vacío.")
            return
        params = []
        for rv in self.param_rows:
            params.append({k: v.get() for k, v in rv.items()})
        config["micro_formats"][name] = {
            "tipo":        self.tipo_var.get(),
            "descripcion": self.desc_var.get(),
            "parametros":  params
        }
        save_config(config)
        self._refresh_list()
        self.slbl['text'] = f"✔ '{name}' guardado."

    def _new(self):
        self.nombre_var.set("Nuevo Formato")
        self.tipo_var.set("simple")
        self.desc_var.set("")
        for w in self.psf.winfo_children(): w.destroy()
        self.param_rows = []

    def _move_up(self):
        sel = self.listbox.curselection()
        if not sel or sel[0] == 0: return
        idx = sel[0]
        keys = list(config["micro_formats"].keys())
        keys[idx-1], keys[idx] = keys[idx], keys[idx-1]
        config["micro_formats"] = {k: config["micro_formats"][k] for k in keys}
        save_config(config)
        self._refresh_list()
        self.listbox.selection_set(idx - 1)
        self.listbox.see(idx - 1)

    def _move_down(self):
        sel = self.listbox.curselection()
        keys = list(config["micro_formats"].keys())
        if not sel or sel[0] >= len(keys) - 1: return
        idx = sel[0]
        keys[idx], keys[idx+1] = keys[idx+1], keys[idx]
        config["micro_formats"] = {k: config["micro_formats"][k] for k in keys}
        save_config(config)
        self._refresh_list()
        self.listbox.selection_set(idx + 1)
        self.listbox.see(idx + 1)

    def _delete(self):
        sel = self.listbox.curselection()
        if not sel: return
        name = self.listbox.get(sel[0])
        if messagebox.askyesno("Confirmar", f"Eliminar '{name}'?"):
            config["micro_formats"].pop(name, None)
            save_config(config)
            self._refresh_list()

# ============================================================
# EDITOR DE FORMATOS DE CALIDAD
# ============================================================

class QualityFormatEditor(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("Editor de Formatos de Calidad (Defectos)")
        self.geometry("860x580")
        self.resizable(True, True)
        self.grab_set()
        self._build_ui()

    def _build_ui(self):
        main = ttk.PanedWindow(self, orient=tk.HORIZONTAL)
        main.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        left = ttk.Frame(main, width=200)
        main.add(left, weight=1)
        ttk.Label(left, text="Formatos:", font=("Helvetica", 10, "bold")).pack(anchor=tk.W)
        lb_frame = ttk.Frame(left)
        lb_frame.pack(fill=tk.BOTH, expand=True)
        self.listbox = tk.Listbox(lb_frame, width=22, height=18, selectmode=tk.SINGLE)
        self.listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        lb_scroll = ttk.Scrollbar(lb_frame, orient="vertical", command=self.listbox.yview)
        lb_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.listbox.config(yscrollcommand=lb_scroll.set)
        self.listbox.bind("<<ListboxSelect>>", self._on_select)
        bf = ttk.Frame(left)
        bf.pack(fill=tk.X, pady=5)
        ttk.Button(bf, text="+ Nuevo",  command=self._new,   width=7).pack(side=tk.LEFT, padx=1)
        ttk.Button(bf, text="Borrar",   command=self._delete,width=7).pack(side=tk.LEFT, padx=1)
        ttk.Button(bf, text="Subir",    command=self._move_up,  width=6).pack(side=tk.LEFT, padx=1)
        ttk.Button(bf, text="Bajar",    command=self._move_down,width=6).pack(side=tk.LEFT, padx=1)

        right = ttk.Frame(main)
        main.add(right, weight=4)

        top = ttk.Frame(right)
        top.pack(fill=tk.X, pady=5)
        ttk.Label(top, text="Nombre:").grid(row=0, column=0, sticky=tk.W, padx=5)
        self.nombre_var = tk.StringVar()
        ttk.Entry(top, textvariable=self.nombre_var, width=24).grid(row=0, column=1, padx=5)
        ttk.Label(top, text="Columnas:").grid(row=0, column=2, sticky=tk.W, padx=5)
        self.cols_var = tk.StringVar(value="2")
        ttk.Combobox(top, textvariable=self.cols_var,
                     values=["2", "3"], state="readonly", width=6).grid(row=0, column=3, padx=5)
        ttk.Label(top, text="Descripción:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        self.desc_var = tk.StringVar()
        ttk.Entry(top, textvariable=self.desc_var, width=55).grid(row=1, column=1, columnspan=3, padx=5)

        hdr = ttk.Frame(right)
        hdr.pack(fill=tk.X, padx=5, pady=(4,0))
        ttk.Label(hdr, text="Nombre parámetro", font=("Helvetica", 8, "bold"),
                  width=28).grid(row=0, column=0, padx=2)
        ttk.Label(hdr, text="Limit (solo col 3)", font=("Helvetica", 8, "bold"),
                  width=16).grid(row=0, column=1, padx=2)
        ttk.Label(hdr, text="Defecto", font=("Helvetica", 8, "bold"),
                  width=10).grid(row=0, column=2, padx=2)

        cf = ttk.Frame(right)
        cf.pack(fill=tk.BOTH, expand=True, padx=5)
        self.pcanvas = tk.Canvas(cf, highlightthickness=0)
        sb = ttk.Scrollbar(cf, orient="vertical", command=self.pcanvas.yview)
        self.psf = ttk.Frame(self.pcanvas)
        self.psf.bind("<Configure>",
                      lambda e: self.pcanvas.configure(scrollregion=self.pcanvas.bbox("all")))
        self.pcanvas.create_window((0, 0), window=self.psf, anchor="nw")
        self.pcanvas.configure(yscrollcommand=sb.set)
        self.pcanvas.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        self.param_rows = []

        pb = ttk.Frame(right)
        pb.pack(fill=tk.X, padx=5, pady=3)
        ttk.Button(pb, text="+ Agregar parámetro", command=self._add_row).pack(side=tk.LEFT)

        sf2 = ttk.Frame(right)
        sf2.pack(fill=tk.X, padx=5, pady=8)
        ttk.Button(sf2, text="💾 Guardar formato", command=self._save).pack(side=tk.RIGHT)
        self.slbl = ttk.Label(sf2, text="", foreground="green")
        self.slbl.pack(side=tk.LEFT)

        self._refresh_list()

    def _refresh_list(self):
        self.listbox.delete(0, tk.END)
        for n in config["quality_formats"]:
            self.listbox.insert(tk.END, n)

    def _on_select(self, event):
        sel = self.listbox.curselection()
        if not sel: return
        name = self.listbox.get(sel[0])
        fmt  = config["quality_formats"].get(name, {})
        self.nombre_var.set(name)
        self.cols_var.set(str(fmt.get("columnas", 2)))
        self.desc_var.set(fmt.get("descripcion", ""))
        for w in self.psf.winfo_children(): w.destroy()
        self.param_rows = []
        for p in fmt.get("parametros", []):
            self._add_row(p)

    def _add_row(self, param=None):
        i = len(self.param_rows)
        rv = {
            "nombre":  tk.StringVar(value=param.get("nombre",  "") if param else ""),
            "limit":   tk.StringVar(value=param.get("limit",   "") if param else ""),
            "defecto": tk.StringVar(value=param.get("defecto", "") if param else ""),
        }
        ttk.Entry(self.psf, textvariable=rv["nombre"],  width=28).grid(row=i, column=0, padx=2, pady=2)
        ttk.Entry(self.psf, textvariable=rv["limit"],   width=16).grid(row=i, column=1, padx=2, pady=2)
        ttk.Entry(self.psf, textvariable=rv["defecto"], width=10).grid(row=i, column=2, padx=2, pady=2)
        ttk.Button(self.psf, text="✕", width=2,
                   command=lambda r=rv: self._remove_row(r)).grid(row=i, column=3, padx=2)
        self.param_rows.append(rv)

    def _remove_row(self, rv):
        self.param_rows.remove(rv)
        for w in self.psf.winfo_children(): w.destroy()
        temp = list(self.param_rows)
        self.param_rows = []
        for p in temp:
            self._add_row({k: v.get() for k, v in p.items()})

    def _new(self):
        self.nombre_var.set("Nuevo Formato")
        self.cols_var.set("2")
        self.desc_var.set("")
        for w in self.psf.winfo_children(): w.destroy()
        self.param_rows = []

    def _move_up(self):
        sel = self.listbox.curselection()
        if not sel or sel[0] == 0: return
        idx = sel[0]
        keys = list(config["quality_formats"].keys())
        keys[idx-1], keys[idx] = keys[idx], keys[idx-1]
        config["quality_formats"] = {k: config["quality_formats"][k] for k in keys}
        save_config(config)
        self._refresh_list()
        self.listbox.selection_set(idx - 1)
        self.listbox.see(idx - 1)

    def _move_down(self):
        sel = self.listbox.curselection()
        keys = list(config["quality_formats"].keys())
        if not sel or sel[0] >= len(keys) - 1: return
        idx = sel[0]
        keys[idx], keys[idx+1] = keys[idx+1], keys[idx]
        config["quality_formats"] = {k: config["quality_formats"][k] for k in keys}
        save_config(config)
        self._refresh_list()
        self.listbox.selection_set(idx + 1)
        self.listbox.see(idx + 1)

    def _delete(self):
        sel = self.listbox.curselection()
        if not sel: return
        name = self.listbox.get(sel[0])
        if messagebox.askyesno("Eliminar", f"Eliminar formato '{name}'?"):
            config["quality_formats"].pop(name, None)
            save_config(config)
            self._refresh_list()

    def _save(self):
        name = self.nombre_var.get().strip()
        if not name:
            messagebox.showwarning("Error", "El nombre no puede estar vacío.")
            return
        params = []
        for rv in self.param_rows:
            p = {k: v.get() for k, v in rv.items()}
            if p["nombre"].strip():
                params.append(p)
        config["quality_formats"][name] = {
            "columnas":    int(self.cols_var.get()),
            "descripcion": self.desc_var.get(),
            "parametros":  params,
        }
        save_config(config)
        self._refresh_list()
        self.slbl.config(text=f"✔ Formato '{name}' guardado.")

# ============================================================
# INTERFAZ PRINCIPAL
# ============================================================

# Mapa cliente → formato microbiologico
# Cada entrada: (keyword_en_cliente_o_producto, formato)
def _detect_micro_format(cliente, producto=""):
    """Detecta formato micro por nombre de cliente y/o marca en el producto.
    Lee los mapas desde config para que sean editables sin tocar el codigo."""
    cliente_low  = (cliente  or '').lower()
    producto_low = (producto or '').lower()

    def _resolve(fmt_name):
        for cfg_name in config.get('micro_formats', {}):
            if cfg_name.lower().replace(' ','').startswith(
                    fmt_name.lower().replace(' ','')[:6]):
                return cfg_name
        return fmt_name

    for entry in config.get('client_map', DEFAULT_CLIENT_MAP):
        if entry['keyword'].lower() in cliente_low:
            return _resolve(entry['formato'])

    for entry in config.get('brand_map', DEFAULT_BRAND_MAP):
        if entry['keyword'].lower() in producto_low:
            return _resolve(entry['formato'])

    micro_keys = list(config.get('micro_formats', {}).keys())
    return micro_keys[0] if micro_keys else ''

# Mapa de palabras clave para auto-deteccion de formato de calidad
QUALITY_KEYWORDS = {
    "Arandano":           ["blueberr", "arandano"],
    "Frutilla":           ["strawberr", "frutilla"],
    "Cereza":             ["cherr", "cereza"],
    "Mora":               ["blackberr", "mora"],
    "Mixes / Berries":    ["berry blend", "mixed berr", "triple berr", "berr blend"],
    "Frambuesa / Crumble":["raspberr crumble", "frambuesa crumble"],
    "Frambuesa":          ["raspberr", "frambuesa"],
    "Mango":              ["mango"],
    "Pina":               ["pineapple", "pina"],
    "Dragon Fruit":       ["dragon fruit", "pitahaya", "pitaya"],
    "Palta":              ["avocado", "palta"],
    "Banana":             ["banana", "platano"],
    "Durazno":            ["peach", "durazno", "nectarin"],
    "Arilos":             ["aril", "pomegranate", "granada"],
}

def _detect_quality_format(producto):
    prod_low = producto.lower()
    for fmt_name, kws in sorted(QUALITY_KEYWORDS.items(),
                                 key=lambda x: -max(len(k) for k in x[1])):
        for kw in kws:
            if kw in prod_low:
                # Map back to actual config key
                for cfg_name in config.get('quality_formats', {}):
                    if cfg_name.lower().replace(' ','').startswith(fmt_name.lower().replace(' ','')[:6]):
                        return cfg_name
                return fmt_name
    return list(config.get('quality_formats', {}).keys())[0] if config.get('quality_formats') else ''

class COAGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Generador de Certificados de Análisis (COA)")
        self.root.geometry("860x800")
        self.root.minsize(700, 600)

        self.pdf_path            = tk.StringVar()
        self.output_folder       = tk.StringVar(value=config.get("output_folder", ""))
        self.last_generated_file = None
        self.all_data            = {}
        self.product_widgets     = {}
        self.full_pdf_path       = ""
        # Guarda los resultados micro ingresados para poder revisarlos después
        self._micro_results      = {}
        self._micro_formatos     = {}
        self._productos_lotes    = {}
        self._productos_agg      = {}

        style = ttk.Style()
        try:
            style.theme_use("clam")
        except Exception:
            pass

        # ── Paleta Frutiger Aero ────────────────────────────────────────
        # Fondos celeste suave, sin brillo excesivo
        BG        = "#C8DFF0"   # celeste suave, fondo principal
        BG2       = "#D8EAF5"   # celeste un poco más claro, frames interiores
        PANEL     = "#E4F1FA"   # paneles / LabelFrames
        ACCENT    = "#3A7FC1"   # azul Frutiger medio
        ACCENT2   = "#255F96"   # azul más oscuro para hover
        ACCENT_LT = "#B0D0EB"   # azul muy claro para tab inactiva
        FG        = "#0D2137"   # texto principal — azul muy oscuro
        FG_MID    = "#2A4A6B"   # texto secundario
        FG_DIS    = "#3A5A7A"   # texto desactivado — oscuro y legible
        BORDER    = "#8AB4D4"   # bordes
        SUCCESS   = "#2E7D52"

        self.root.configure(bg=BG)

        style.configure(".",
            background=BG, foreground=FG,
            font=("Segoe UI", 9), borderwidth=0)

        style.configure("TFrame",   background=BG, padding=0)
        style.configure("TLabel",   background=BG, foreground=FG,
                        font=("Segoe UI", 9), padding=(2,2))

        # Notebook — pestaña activa más grande/elevada que las inactivas
        style.configure("TNotebook", background=BG, tabmargins=[2, 4, 0, 0])
        style.configure("TNotebook.Tab",
            background=ACCENT_LT, foreground=FG_MID,
            font=("Segoe UI", 9),
            padding=(14, 5))
        style.map("TNotebook.Tab",
            background=[("selected", PANEL)],
            foreground=[("selected", ACCENT2)],
            font=[("selected", ("Segoe UI", 10, "bold"))],
            padding=[("selected", (16, 7))])

        style.configure("TLabelframe",
            background=PANEL, relief="groove",
            bordercolor=BORDER, borderwidth=1)
        style.configure("TLabelframe.Label",
            background=PANEL, foreground=ACCENT2,
            font=("Segoe UI", 9, "bold"), padding=(4,2))

        # Botones normales
        style.configure("TButton",
            background=ACCENT, foreground="white",
            font=("Segoe UI", 9, "bold"),
            padding=(10, 5), relief="flat", borderwidth=0)
        style.map("TButton",
            background=[("active", ACCENT2), ("disabled", ACCENT_LT)],
            foreground=[("active", "white"), ("disabled", FG_DIS)])

        style.configure("TEntry",
            fieldbackground="white", foreground=FG,
            bordercolor=BORDER, lightcolor=BORDER,
            darkcolor=BORDER, insertcolor=FG,
            padding=(6, 4), relief="flat")
        style.map("TEntry",
            fieldbackground=[("focus", "#EDF6FF")],
            bordercolor=[("focus", ACCENT)])

        style.configure("TCombobox",
            fieldbackground="white", foreground=FG,
            padding=(6, 4), relief="flat")
        style.map("TCombobox",
            fieldbackground=[("readonly", "white")],
            foreground=[("disabled", FG_DIS)])

        # Botón grande de generar
        style.configure("Accent.TButton",
            background=ACCENT, foreground="white",
            font=("Segoe UI", 10, "bold"), padding=(14, 7))
        style.map("Accent.TButton",
            background=[("active", ACCENT2), ("disabled", ACCENT_LT)],
            foreground=[("disabled", FG_DIS)])

        style.configure("Success.TButton",
            background=SUCCESS, foreground="white",
            font=("Segoe UI", 9, "bold"), padding=(10, 5))
        style.map("Success.TButton",
            background=[("active", "#1C5C3A")],
            foreground=[("disabled", FG_DIS)])

        style.configure("Small.TButton",
            background="#4A8EC2", foreground="white",
            font=("Segoe UI", 8), padding=(6, 3))
        style.map("Small.TButton",
            background=[("active", "#C0392B"), ("disabled", ACCENT_LT)],
            foreground=[("active", "white"), ("disabled", FG_DIS)])

        style.configure("TSeparator", background=BORDER)
        style.configure("TScrollbar",
            background=ACCENT_LT, troughcolor=BG2,
            arrowcolor=ACCENT2, borderwidth=0)

        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=8, pady=(8,0))
        self.tab_gen    = ttk.Frame(self.notebook)
        self.tab_config = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_gen,    text="  📋 Generador  ")
        self.notebook.add(self.tab_config, text="  ⚙ Configuración  ")
        # Botón Acerca de — esquina superior derecha
        about_btn = tk.Button(root, text="ℹ  Acerca de",
            bg="#255F96", fg="white", relief="flat",
            font=("Segoe UI", 8), cursor="hand2",
            command=self._open_about, padx=8, pady=3)
        about_btn.place(relx=1.0, rely=0.0, anchor="ne", x=-4, y=4)

        self._build_generator_tab()
        self._build_config_tab()

        # Barra de progreso (encima del status bar)
        self.progress_bar = ttk.Progressbar(root, orient="horizontal",
            mode="determinate", length=100)
        # No se hace pack aquí — se muestra solo al generar

        # Status bar moderna — dos secciones
        status_bar = tk.Frame(root, bg="#255F96", height=28)
        status_bar.pack(fill=tk.X, side=tk.BOTTOM)
        status_bar.pack_propagate(False)

        self.status_label = tk.Label(status_bar, text="  Listo para empezar.",
            bg="#255F96", fg="white",
            font=("Segoe UI", 9), anchor=tk.W)
        self.status_label.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=6)

        # Info embarque/PO siempre visible a la derecha
        self.work_label = tk.Label(status_bar, text="",
            bg="#1A4F7A", fg="#B8D9F5",
            font=("Segoe UI", 9, "bold"), anchor=tk.E, padx=10)
        self.work_label.pack(side=tk.RIGHT)

    def _open_about(self):
        """Ventana Acerca de."""
        win = tk.Toplevel(self.root)
        win.title("Acerca de")
        win.geometry("420x250")
        win.resizable(False, False)
        win.configure(bg="#C8DFF0")
        win.grab_set()
        win.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width()  - 420) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 250) // 2
        win.geometry(f"+{x}+{y}")

        tk.Label(win, text="Generador de COAs",
            bg="#C8DFF0", fg="#0D2137",
            font=("Segoe UI", 16, "bold")).pack(pady=(28, 8))
        sep = tk.Frame(win, bg="#8AB4D4", height=1)
        sep.pack(fill=tk.X, padx=40, pady=(0, 18))
        tk.Label(win,
            text="Desarrollado por José Ibarra\nControl de Calidad — Exportación",
            bg="#C8DFF0", fg="#2A4A6B",
            font=("Segoe UI", 10), justify="center").pack(pady=(0, 10))
        tk.Label(win,
            text="Herramienta interna para generación de COAs\ny registro operativo.",
            bg="#C8DFF0", fg="#255F96",
            font=("Segoe UI", 9), justify="center").pack(pady=(0, 18))
        ttk.Button(win, text="Cerrar", command=win.destroy).pack(pady=12)

    def _apply_arrow_nav_main(self):
        """Registra navegación con flechas en los campos editables de la ventana principal."""
        entries = []
        def _collect(w):
            if isinstance(w, (ttk.Entry, tk.Entry)):
                try:
                    if w.winfo_viewable():
                        entries.append(w)
                except Exception:
                    pass
            for child in w.winfo_children():
                _collect(child)
        _collect(self.root)
        if entries:
            register_arrow_nav(entries)
        # Re-registrar cada vez que se cargue un PDF (campos de calidad cambian)
        self.root.after(2000, self._apply_arrow_nav_main)

    def _build_generator_tab(self):
        # ── Canvas + scrollbar vertical ────────────────────────────────
        outer = ttk.Frame(self.tab_gen)
        outer.pack(fill=tk.BOTH, expand=True)

        vsb = ttk.Scrollbar(outer, orient="vertical")
        self._gen_canvas = tk.Canvas(outer, yscrollcommand=vsb.set,
                                     highlightthickness=0, bg="#C8DFF0",
                                     bd=0)
        vsb.config(command=self._gen_canvas.yview)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        self._gen_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        main = ttk.Frame(self._gen_canvas)
        self._gen_canvas_win = self._gen_canvas.create_window(
            (0, 0), window=main, anchor="nw")

        def _on_frame_configure(e):
            self._gen_canvas.configure(scrollregion=self._gen_canvas.bbox("all"))
        def _on_canvas_configure(e):
            self._gen_canvas.itemconfig(self._gen_canvas_win, width=e.width)
        main.bind("<Configure>", _on_frame_configure)
        self._gen_canvas.bind("<Configure>", _on_canvas_configure)

        def _on_mousewheel(e):
            w = e.widget
            while w:
                if isinstance(w, tk.Canvas):
                    w.yview_scroll(int(-1*(e.delta/120)), "units")
                    return
                try:
                    w = w.master
                except Exception:
                    break
        self._gen_canvas.bind_all("<MouseWheel>", _on_mousewheel)

        PAD = dict(padx=14, pady=5)

        # ── Paso 1 ──────────────────────────────────────────────────────
        self._step_titles = {}
        self._step_frames = {}
        s1 = ttk.LabelFrame(main, text="  1 · Packing List", padding=(12,8))
        self._step_frames[1] = s1
        s1.pack(fill=tk.X, **PAD)
        row1 = ttk.Frame(s1)
        row1.pack(fill=tk.X)
        ttk.Button(row1, text="📂  Seleccionar PDF",
                   command=self.select_pdf).pack(side=tk.LEFT, padx=(0,6))
        self.view_pdf_button = ttk.Button(row1, text="👁  Ver PDF",
                   command=self.open_selected_pdf, state=tk.DISABLED)
        self.view_pdf_button.pack(side=tk.LEFT, padx=(0,10))
        self.pdf_label = ttk.Label(row1, textvariable=self.pdf_path,
                   foreground="#255F96", font=("Segoe UI", 9, "italic"))
        self.pdf_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
        # Botón limpiar — pequeño, a la derecha
        ttk.Button(row1, text="✖  Nueva carga",
                   command=self.reset_all,
                   style="Small.TButton").pack(side=tk.RIGHT, padx=(6,0))

        # ── Paso 2 ──────────────────────────────────────────────────────
        s2 = ttk.LabelFrame(main, text="  2 · Cargar Datos", padding=(12,8))
        self._step_frames[2] = s2
        s2.pack(fill=tk.X, **PAD)
        row2 = ttk.Frame(s2)
        row2.pack(fill=tk.X)
        self.process_button = ttk.Button(row2, text="⬇  Cargar Datos del PDF",
                   command=self.process_pdf, state=tk.DISABLED)
        self.process_button.pack(side=tk.LEFT, padx=(0,12))
        self.summary_text = tk.Text(s2, height=4, state=tk.DISABLED,
                   background="#D8EAF5", relief=tk.FLAT,
                   font=("Consolas", 9), fg="#0D2137",
                   padx=8, pady=6, wrap=tk.NONE)
        self.summary_text.pack(fill=tk.X, pady=(8,0))

        # ── Paso 2.5 ────────────────────────────────────────────────────
        self.edit_embarque   = tk.StringVar()
        self.edit_cliente    = tk.StringVar()
        self.edit_po         = tk.StringVar()
        self.edit_contenedor = tk.StringVar()
        self.edit_sello      = tk.StringVar()
        self.edit_recorder   = tk.StringVar()
        self.edit_palletized = tk.StringVar(value="pallet")

        self.s25 = ttk.LabelFrame(main, text="  2.5 · Verificar y Corregir Datos", padding=(12,8))
        # NO se hace pack aquí — se muestra en process_pdf

        grid25 = ttk.Frame(self.s25)
        grid25.pack(fill=tk.X, pady=(0,4))
        campos = [
            ("Embarque:",      self.edit_embarque),
            ("Cliente:",       self.edit_cliente),
            ("PO Number:",     self.edit_po),
            ("Contenedor:",    self.edit_contenedor),
            ("Sello:",         self.edit_sello),
            ("Ryan Recorder:", self.edit_recorder),
        ]
        self._edit_entries_25 = []
        for i, (lbl, var) in enumerate(campos):
            col = (i % 3) * 2
            row = i // 3
            ttk.Label(grid25, text=lbl, anchor="e",
                      font=("Segoe UI", 9), foreground="#2A4A6B").grid(
                row=row, column=col, sticky=tk.E, padx=(8,4), pady=4)
            e = ttk.Entry(grid25, textvariable=var, width=22)
            e.grid(row=row, column=col+1, sticky=tk.EW, padx=(0,10), pady=4)
            self._edit_entries_25.append(e)
        for c in range(6):
            grid25.columnconfigure(c, weight=1 if c % 2 == 1 else 0)

        pal_row = ttk.Frame(self.s25)
        pal_row.pack(anchor=tk.W, padx=6, pady=(0,6))
        ttk.Label(pal_row, text="Tipo de carga:",
                  foreground="#2A4A6B").pack(side=tk.LEFT, padx=(0,6))
        ttk.Combobox(pal_row, textvariable=self.edit_palletized,
                     values=["pallet", "floor", "slipsheet"],
                     state="readonly", width=14).pack(side=tk.LEFT)

        self.brix_ph_frame = ttk.LabelFrame(self.s25,
                     text="  Brix, pH y Calidad por Producto", padding=(8,6))
        self.brix_ph_frame.pack(fill=tk.X, padx=6, pady=(4,4))
        self.brix_ph_vars    = {}
        self.quality_formato = {}
        self.quality_vars    = {}

        # Registrar navegación con flechas en los campos de Paso 2.5
        # (se llama después del pack vía after)
        self.root.after(200, self._apply_arrow_nav_main)

        # ── Paso 3 ──────────────────────────────────────────────────────
        s3 = ttk.LabelFrame(main, text="  3 · Plantillas por Producto", padding=(12,8))
        self._step_frames[3] = s3
        s3.pack(fill=tk.X, **PAD)
        self.products_frame = ttk.Frame(s3)
        self.products_frame.pack(fill=tk.X)

        # ── Paso 4 ──────────────────────────────────────────────────────
        s4 = ttk.LabelFrame(main, text="  4 · Carpeta de Destino", padding=(12,8))
        self._step_frames[4] = s4
        s4.pack(fill=tk.X, **PAD)
        row4 = ttk.Frame(s4)
        row4.pack(fill=tk.X)
        ttk.Button(row4, text="📁  Seleccionar Carpeta",
                   command=self.select_output_folder).pack(side=tk.LEFT, padx=(0,10))
        ttk.Label(row4, textvariable=self.output_folder,
                  foreground="#255F96",
                  font=("Segoe UI", 9, "italic")).pack(side=tk.LEFT, fill=tk.X, expand=True)

        # ── Paso 5 ──────────────────────────────────────────────────────
        s5 = ttk.LabelFrame(main, text="  5 · Microbiología", padding=(12,8))
        self._step_frames[5] = s5
        s5.pack(fill=tk.X, **PAD)
        row5 = ttk.Frame(s5)
        row5.pack(fill=tk.X)
        self.micro_button = ttk.Button(row5, text="🔬  Ingresar / Revisar Microbiología",
                   command=self.open_micro_dialog, state=tk.DISABLED)
        self.micro_button.pack(side=tk.LEFT, padx=(0,12))
        self.micro_status_label = ttk.Label(row5, text="Sin resultados ingresados.",
                   foreground="#2A4A6B", font=("Segoe UI", 9, "italic"))
        self.micro_status_label.pack(side=tk.LEFT)

        # ── Paso 6 ──────────────────────────────────────────────────────
        s6 = ttk.LabelFrame(main, text="  6 · Generar Archivos", padding=(12,10))
        self._step_frames[6] = s6
        s6.pack(fill=tk.X, **PAD)
        row6a = ttk.Frame(s6)
        row6a.pack(fill=tk.X, pady=(0,6))
        self.generate_button = ttk.Button(row6a, text="✅  Generar COAs",
                   command=self.generate_coas, state=tk.DISABLED,
                   style="Accent.TButton")
        self.generate_button.pack(side=tk.LEFT, padx=(0,8))
        self.open_file_button   = ttk.Button(row6a, text="📄 Abrir COA",
                   command=self.open_last_file)
        self.open_folder_button = ttk.Button(row6a, text="📂 Abrir Carpeta",
                   command=self.open_output_folder)
        row6b = ttk.Frame(s6)
        row6b.pack(fill=tk.X)
        ttk.Button(row6b, text="📋 Ver Registro",
                   command=self.open_registro).pack(side=tk.LEFT, padx=(0,6))

    def _build_config_tab(self):
        cfg_nb = ttk.Notebook(self.tab_config)
        cfg_nb.pack(fill=tk.BOTH, expand=True, padx=6, pady=6)
        tab_rutas = ttk.Frame(cfg_nb)
        tab_mapas = ttk.Frame(cfg_nb)
        tab_log   = ttk.Frame(cfg_nb)
        cfg_nb.add(tab_rutas, text="  Rutas y Formatos  ")
        cfg_nb.add(tab_mapas, text="  Clientes / Marcas  ")
        cfg_nb.add(tab_log,   text="  Log de Errores  ")
        self._build_rutas_tab(tab_rutas)
        self._build_mapas_tab(tab_mapas)
        self._build_log_tab(tab_log)

    def _build_rutas_tab(self, parent):
        outer = ttk.Frame(parent)
        outer.pack(fill=tk.BOTH, expand=True, padx=20, pady=16)
        ttk.Label(outer, text="Rutas por defecto",
                  font=("Segoe UI", 11, "bold"),
                  foreground="#255F96").pack(anchor=tk.W, pady=(0,12))
        self.cfg_pdf_folder      = tk.StringVar(value=config.get("pdf_folder", ""))
        self.cfg_template_folder = tk.StringVar(value=config.get("template_folder", ""))
        self.cfg_output_folder   = tk.StringVar(value=config.get("output_folder", ""))
        rutas = [
            ("Packing Lists (PDFs):", self.cfg_pdf_folder),
            ("Plantillas (Templates):", self.cfg_template_folder),
            ("Destino (COAs):",         self.cfg_output_folder),
        ]
        frame = ttk.Frame(outer)
        frame.pack(fill=tk.X)
        frame.columnconfigure(1, weight=1)
        for i, (lbl, var) in enumerate(rutas):
            ttk.Label(frame, text=lbl, foreground="#2A4A6B",
                      font=("Segoe UI", 9)).grid(
                row=i, column=0, sticky=tk.W, pady=6, padx=(0,10))
            ttk.Entry(frame, textvariable=var).grid(
                row=i, column=1, sticky=tk.EW, padx=(0,8), pady=6)
            ttk.Button(frame, text="Examinar",
                       command=lambda v=var: self._browse_folder(v)).grid(
                row=i, column=2, pady=6)
        ttk.Separator(outer, orient="horizontal").pack(fill=tk.X, pady=12)
        fmt_f = ttk.Frame(outer)
        fmt_f.pack(fill=tk.X)
        ttk.Label(fmt_f, text="Formatos Microbiologicos:",
                  font=("Segoe UI", 9, "bold"), foreground="#2A4A6B").grid(
            row=0, column=0, sticky=tk.W, pady=5)
        ttk.Button(fmt_f, text="Editar Formatos Microbiologicos",
                   command=lambda: MicroFormatEditor(self.root)).grid(
            row=0, column=1, sticky=tk.W, padx=10)
        ttk.Label(fmt_f, text="Formatos de Calidad (Defectos):",
                  font=("Segoe UI", 9, "bold"), foreground="#2A4A6B").grid(
            row=1, column=0, sticky=tk.W, pady=5)
        ttk.Button(fmt_f, text="Editar Formatos de Calidad",
                   command=lambda: QualityFormatEditor(self.root)).grid(
            row=1, column=1, sticky=tk.W, padx=10)
        ttk.Separator(outer, orient="horizontal").pack(fill=tk.X, pady=12)
        bot_row = ttk.Frame(outer)
        bot_row.pack(anchor=tk.W)
        ttk.Button(bot_row, text="Guardar Configuracion",
                   command=self.save_configuration,
                   style="Accent.TButton").pack(side=tk.LEFT, padx=(0,10))
        self.config_status_label = ttk.Label(bot_row, text="", foreground="#1E7C42",
                                             font=("Segoe UI", 9, "bold"))
        self.config_status_label.pack(side=tk.LEFT)

    def _build_mapas_tab(self, parent):
        """Pestaña editable de mapas cliente-formato y marca-formato."""
        outer = ttk.Frame(parent)
        outer.pack(fill=tk.BOTH, expand=True, padx=12, pady=10)
        outer.columnconfigure(0, weight=1)
        outer.columnconfigure(1, weight=1)

        def _make_panel(col, titulo, data_key, default_data):
            lf = ttk.LabelFrame(outer, text=titulo, padding=(10, 8))
            lf.grid(row=0, column=col, sticky="nsew", padx=(0, 6) if col == 0 else (6, 0))
            lf.columnconfigure(0, weight=1)
            lf.columnconfigure(1, weight=1)
            ttk.Label(lf, text="Palabra clave", font=("Segoe UI", 8, "bold"),
                      foreground="#255F96").grid(row=0, column=0, sticky=tk.W, padx=4, pady=2)
            ttk.Label(lf, text="Formato micro", font=("Segoe UI", 8, "bold"),
                      foreground="#255F96").grid(row=0, column=1, sticky=tk.W, padx=4, pady=2)
            rows_f = ttk.Frame(lf)
            rows_f.grid(row=1, column=0, columnspan=3, sticky="ew")
            rows_f.columnconfigure(0, weight=1)
            rows_f.columnconfigure(1, weight=1)
            entry_rows = []
            fmt_keys = list(config.get("micro_formats", {}).keys())

            def _add_row(kw="", fmt="", rows_f=rows_f, entry_rows=entry_rows, fmt_keys=fmt_keys):
                r = len(entry_rows)
                kw_v  = tk.StringVar(value=kw)
                fmt_v = tk.StringVar(value=fmt)
                kw_e  = ttk.Entry(rows_f, textvariable=kw_v, width=18)
                fmt_c = ttk.Combobox(rows_f, textvariable=fmt_v,
                                     values=fmt_keys, state="readonly", width=18)
                del_b = ttk.Button(rows_f, text="x", width=2,
                                   command=lambda i=r: _del_row(i, entry_rows))
                kw_e.grid( row=r, column=0, padx=3, pady=2, sticky="ew")
                fmt_c.grid(row=r, column=1, padx=3, pady=2, sticky="ew")
                del_b.grid(row=r, column=2, padx=2, pady=2)
                entry_rows.append([kw_v, fmt_v, kw_e, fmt_c, del_b])

            def _del_row(idx, entry_rows):
                if idx < len(entry_rows) and entry_rows[idx]:
                    for w in entry_rows[idx][2:]:
                        w.destroy()
                    entry_rows[idx] = None

            def _save(data_key=data_key, entry_rows=entry_rows):
                data = [{"keyword": r[0].get().strip().lower(),
                         "formato": r[1].get().strip()}
                        for r in entry_rows if r and r[0].get().strip()]
                config[data_key] = data
                save_config(config)
                self.config_status_label["text"] = "Guardado."
                self.root.after(3000, lambda: self.config_status_label.config(text=""))

            for entry in config.get(data_key, default_data):
                _add_row(entry.get("keyword", ""), entry.get("formato", ""))

            ttk.Button(lf, text="+ Agregar", command=_add_row,
                       style="Small.TButton").grid(row=2, column=0, sticky=tk.W, pady=(6, 2))
            ttk.Button(lf, text="Guardar", command=_save,
                       style="Small.TButton").grid(row=2, column=1, sticky=tk.W, pady=(6, 2))

        _make_panel(0, "  Clientes", "client_map", DEFAULT_CLIENT_MAP)
        _make_panel(1, "  Marcas en Producto", "brand_map", DEFAULT_BRAND_MAP)

    def _build_log_tab(self, parent):
        """Pestaña con el contenido del error_log.txt."""
        outer = ttk.Frame(parent)
        outer.pack(fill=tk.BOTH, expand=True, padx=12, pady=10)
        btn_row = ttk.Frame(outer)
        btn_row.pack(fill=tk.X, pady=(0, 6))
        ttk.Label(btn_row, text="Ultimos errores registrados:",
                  font=("Segoe UI", 9, "bold"), foreground="#255F96").pack(side=tk.LEFT)
        ttk.Button(btn_row, text="Actualizar",
                   command=lambda: self._refresh_log(log_text),
                   style="Small.TButton").pack(side=tk.LEFT, padx=8)
        ttk.Button(btn_row, text="Limpiar log",
                   command=lambda: self._clear_log(log_text),
                   style="Small.TButton").pack(side=tk.LEFT)
        log_frame = ttk.Frame(outer)
        log_frame.pack(fill=tk.BOTH, expand=True)
        vsb = ttk.Scrollbar(log_frame, orient="vertical")
        log_text = tk.Text(log_frame, height=18, state=tk.DISABLED,
                           font=("Consolas", 8), bg="#0D1B2A", fg="#A8D8EA",
                           relief=tk.FLAT, padx=8, pady=6,
                           yscrollcommand=vsb.set, wrap=tk.NONE)
        vsb.config(command=log_text.yview)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self._refresh_log(log_text)

    def _refresh_log(self, text_widget):
        text_widget.config(state=tk.NORMAL)
        text_widget.delete(1.0, tk.END)
        try:
            log_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "error_log.txt")
            if os.path.exists(log_path):
                with open(log_path, "r", encoding="utf-8", errors="replace") as f:
                    lineas = f.read().splitlines()[-200:]
                text_widget.insert(tk.END, "\n".join(lineas))
                text_widget.see(tk.END)
            else:
                text_widget.insert(tk.END, "No hay errores registrados.")
        except Exception as e:
            text_widget.insert(tk.END, f"Error leyendo log: {e}")
        text_widget.config(state=tk.DISABLED)

    def _clear_log(self, text_widget):
        if not messagebox.askyesno("Limpiar log", "Borrar el registro de errores?"):
            return
        try:
            log_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "error_log.txt")
            open(log_path, "w").close()
            self._refresh_log(text_widget)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo limpiar: {e}")

    def _browse_folder(self, sv):
        path = filedialog.askdirectory()
        if path: sv.set(path)

    def _check_embarque_duplicado(self, embarque, po=""):
        """Revisa Registro_COAs.xlsx y avisa si el embarque ya fue procesado."""
        try:
            import openpyxl
            rpath = get_registro_path()
            if not os.path.exists(rpath):
                return
            wb = openpyxl.load_workbook(rpath, read_only=True)
            encontrados = []
            for sname in wb.sheetnames:
                ws = wb[sname]
                for row in ws.iter_rows(min_row=2, values_only=True):
                    if row and str(row[1] or "").strip() == embarque.strip():
                        fecha   = str(row[0] or "")
                        cliente = str(row[2] or "")
                        prod    = str(row[4] or "")
                        if not any(e[0]==fecha and e[2]==prod for e in encontrados):
                            encontrados.append((fecha, cliente, prod))
            wb.close()
            if encontrados:
                lineas = ["  \u2022 " + e[2] + " \u2014 " + e[1] + " (" + e[0] + ")"
                          for e in encontrados[:8]]
                detalle = "\n".join(lineas)
                if len(encontrados) > 8:
                    detalle += f"\n  ... y {len(encontrados)-8} mas"
                messagebox.showwarning(
                    "\u26a0 Embarque ya registrado",
                    f"El embarque '{embarque}' ya tiene COAs generados:\n\n"
                    f"{detalle}\n\n"
                    "Puedes continuar, pero verifica que no sea un duplicado."
                )
        except Exception:
            pass

    def save_configuration(self):
        global config
        config["pdf_folder"]      = self.cfg_pdf_folder.get()
        config["template_folder"] = self.cfg_template_folder.get()
        config["output_folder"]   = self.cfg_output_folder.get()
        save_config(config)
        self.output_folder.set(config["output_folder"])
        self.config_status_label['text'] = "✔ Configuración guardada."
        self.check_if_ready_to_generate()

    def open_selected_pdf(self):
        if not self.full_pdf_path: return
        try:
            if sys.platform == "win32":    os.startfile(self.full_pdf_path)
            elif sys.platform == "darwin": subprocess.Popen(["open", self.full_pdf_path])
            else:                          subprocess.Popen(["xdg-open", self.full_pdf_path])
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir el PDF:\n{e}")

    def select_pdf(self):
        path = filedialog.askopenfilename(
            title="Selecciona el Packing List en PDF",
            filetypes=[("PDF files", "*.pdf")],
            initialdir=config.get("pdf_folder", "")
        )
        if path:
            self.pdf_path.set(os.path.basename(path))
            self.full_pdf_path = path
            self.process_button['state']  = tk.NORMAL
            self.view_pdf_button['state'] = tk.NORMAL
            self.status_label['text']     = "PDF seleccionado. Presiona 'Cargar Datos'."
            self._set_step_done(1, True)
            self.clear_products()
        else:
            self.view_pdf_button['state'] = tk.DISABLED

    def open_output_folder(self):
        folder = self.output_folder.get()
        if not folder: return
        try:
            if sys.platform == "win32":    os.startfile(folder)
            elif sys.platform == "darwin": subprocess.Popen(["open", folder])
            else:                          subprocess.Popen(["xdg-open", folder])
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir la carpeta:\n{e}")

    def _build_quality_fields(self, producto):
        """Construye dinámicamente los campos de defectos para el producto."""
        fmt_nombre = self.quality_formato.get(producto, tk.StringVar()).get()
        fmt        = config["quality_formats"].get(fmt_nombre, {})
        params     = fmt.get("parametros", [])
        ncols      = fmt.get("columnas", 2)

        # Buscar el frame de campos del producto
        parent_frame = None
        for w in self.brix_ph_frame.winfo_children():
            if hasattr(w, "_quality_fields_frame"):
                # Verificar que es el producto correcto buscando el título del LabelFrame
                if w.cget("text").startswith(producto[:50]):
                    parent_frame = w._quality_fields_frame
                    break
        if parent_frame is None:
            return

        for w in parent_frame.winfo_children():
            w.destroy()
        self.quality_vars[producto] = {}

        # Encabezado columnas
        header = ttk.Frame(parent_frame)
        header.pack(fill=tk.X, padx=4)
        ttk.Label(header, text="Parámetro", width=30, anchor="w",
                  font=("Helvetica", 8, "bold")).pack(side=tk.LEFT)
        if ncols == 3:
            ttk.Label(header, text="Limit", width=12, anchor="center",
                      font=("Helvetica", 8, "bold")).pack(side=tk.LEFT, padx=(0,4))
        ttk.Label(header, text="Results", width=10, anchor="center",
                  font=("Helvetica", 8, "bold")).pack(side=tk.LEFT)

        ttk.Separator(parent_frame, orient="horizontal").pack(fill=tk.X, padx=4, pady=2)

        for param in params:
            nombre  = param["nombre"]
            defecto = param.get("defecto", "")
            limit   = param.get("limit", "")

            row = ttk.Frame(parent_frame)
            row.pack(fill=tk.X, padx=4, pady=1)
            ttk.Label(row, text=nombre, width=30, anchor="w").pack(side=tk.LEFT)
            if ncols == 3:
                ttk.Label(row, text=limit, width=12, anchor="center",
                          foreground="#2A4A6B").pack(side=tk.LEFT, padx=(0,4))
            var = tk.StringVar(value=defecto)
            ttk.Entry(row, textvariable=var, width=10).pack(side=tk.LEFT)
            self.quality_vars[producto][nombre] = var

        # Botón para agregar parámetro extra
        add_row = ttk.Frame(parent_frame)
        add_row.pack(anchor=tk.W, padx=4, pady=(4,0))
        ttk.Button(add_row, text="+ Agregar parámetro",
                   command=lambda p=producto, pf=parent_frame, n=ncols:
                   self._add_quality_param(p, pf, n)).pack(side=tk.LEFT)

    def _add_quality_param(self, producto, parent_frame, ncols):
        """Agrega un campo de parámetro extra al final de la lista."""
        # Insertar antes del botón de agregar (último widget)
        widgets = parent_frame.winfo_children()
        add_btn_frame = widgets[-1] if widgets else None

        row = ttk.Frame(parent_frame)
        if add_btn_frame:
            row.pack(fill=tk.X, padx=4, pady=1, before=add_btn_frame)
        else:
            row.pack(fill=tk.X, padx=4, pady=1)

        nombre_var = tk.StringVar()
        ttk.Entry(row, textvariable=nombre_var, width=30).pack(side=tk.LEFT)
        if ncols == 3:
            ttk.Label(row, text="", width=12).pack(side=tk.LEFT, padx=(0,4))
        val_var = tk.StringVar()
        ttk.Entry(row, textvariable=val_var, width=10).pack(side=tk.LEFT)

        # Registrar con clave dinámica — se actualiza al escribir el nombre
        key = f"__extra_{len(self.quality_vars[producto])}"
        self.quality_vars[producto][key] = val_var

        def _update_key(*args, k=key):
            new_key = nombre_var.get().strip()
            if new_key and k in self.quality_vars[producto]:
                self.quality_vars[producto][new_key] = self.quality_vars[producto].pop(k)
        nombre_var.trace_add("write", _update_key)


    def agregar_micro_a_coa(self):
        """Permite agregar microbiología a un COA .docx ya generado."""
        # 1. Seleccionar el COA existente
        coa_path = filedialog.askopenfilename(
            title="Selecciona el COA al que agregar microbiología",
            filetypes=[("Word files", "*.docx")],
            initialdir=self.output_folder.get() or config.get("output_folder", "")
        )
        if not coa_path:
            return

        # 2. Pedir nombre de producto y cliente para mostrar la ventana correcta
        info_win = tk.Toplevel(self.root)
        info_win.title("Configurar microbiología")
        info_win.geometry("420x200")
        info_win.grab_set()

        ttk.Label(info_win, text="Producto:").grid(row=0, column=0, padx=10, pady=8, sticky="e")
        prod_var = tk.StringVar(value=os.path.basename(coa_path).replace(".docx",""))
        ttk.Entry(info_win, textvariable=prod_var, width=35).grid(row=0, column=1, padx=10, pady=8)

        ttk.Label(info_win, text="Cliente:").grid(row=1, column=0, padx=10, pady=8, sticky="e")
        cli_var = tk.StringVar(value=self.all_data.get("general_info",{}).get("cliente","") if self.all_data else "")
        ttk.Entry(info_win, textvariable=cli_var, width=35).grid(row=1, column=1, padx=10, pady=8)

        ttk.Label(info_win, text="Lote(s) (separados por coma):").grid(row=2, column=0, padx=10, pady=8, sticky="e")
        lote_var = tk.StringVar()
        ttk.Entry(info_win, textvariable=lote_var, width=35).grid(row=2, column=1, padx=10, pady=8)

        def _continuar():
            producto  = prod_var.get().strip() or "Producto"
            cliente   = cli_var.get().strip()
            lotes_raw = lote_var.get().strip()
            lotes     = [l.strip() for l in lotes_raw.split(",") if l.strip()] if lotes_raw else ["Lote 1"]
            info_win.destroy()

            # 3. Detectar formato micro por cliente
            fmt_auto = _detect_micro_format(cliente, producto)

            # 4. Abrir ventana de microbiología
            micro_history = load_micro_history()
            dialogo = MicrobiologyInputDialog(
                self.root,
                {producto: lotes},
                {producto: fmt_auto},
                cliente,
                micro_history
            )
            self.root.wait_window(dialogo)
            if not dialogo.confirmed:
                return

            micro_results  = dialogo.get_results()
            micro_formatos = dialogo.get_formatos()
            micro_per_lote = micro_results.get(producto, {})
            fmt_nombre     = micro_formatos.get(producto, fmt_auto)
            fmt            = config["micro_formats"].get(fmt_nombre, {})
            tipo_micro     = fmt.get("tipo", "simple")

            # 5. Cargar el COA, llenar micro y guardar
            try:
                doc = Document(coa_path)
                if tipo_micro == "camerican":
                    duplicate_camerican_tables(doc, lotes, micro_per_lote)
                elif tipo_micro == "n_series":
                    duplicate_n_series_tables(doc, lotes, micro_per_lote)
                    micro_tables = find_micro_tables(doc)
                    for _, table, tipo_tabla in micro_tables:
                        if tipo_tabla == "small":
                            fill_micro_small(table, lotes, micro_per_lote)
                            break
                else:
                    micro_tables = find_micro_tables(doc)
                    for _, table, tipo_tabla in micro_tables:
                        if tipo_tabla == "main":
                            fill_micro_simple(doc, table, lotes, micro_per_lote)
                        elif tipo_tabla == "small":
                            fill_micro_small(table, lotes, micro_per_lote)

                doc.save(coa_path)
                for lote_k, lote_vals in micro_per_lote.items():
                    save_micro_history_record(cliente, producto, lote_k, lote_vals, fmt_nombre)
                messagebox.showinfo("Listo",
                    f"Microbiología agregada y guardada en:\n{os.path.basename(coa_path)}")
            except Exception as e:
                logging.exception("Error al agregar micro a COA")
                messagebox.showerror("Error", f"No se pudo actualizar el COA:\n{e}")

        ttk.Button(info_win, text="Continuar →", command=_continuar).grid(
            row=3, column=0, columnspan=2, pady=12)

    def open_registro(self):
        path = get_registro_path()
        if not os.path.exists(path):
            messagebox.showinfo("Registro", "Aún no hay COAs registrados.")
            return
        try:
            if sys.platform == "win32":    os.startfile(path)
            elif sys.platform == "darwin": subprocess.Popen(["open", path])
            else:                          subprocess.Popen(["xdg-open", path])
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir el registro:\n{e}")

    def open_last_file(self):
        if not self.last_generated_file: return
        try:
            if sys.platform == "win32":    os.startfile(self.last_generated_file)
            elif sys.platform == "darwin": subprocess.Popen(["open", self.last_generated_file])
            else:                          subprocess.Popen(["xdg-open", self.last_generated_file])
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir el archivo:\n{e}")

    def select_output_folder(self):
        path = filedialog.askdirectory(
            title="Selecciona la carpeta para guardar los COA",
            initialdir=config.get("output_folder", "")
        )
        if path:
            self._set_step_done(4, True)
        if path:
            self.output_folder.set(path)
            self.check_if_ready_to_generate()

    def process_pdf(self):
        self.clear_products()
        try:
            self.all_data = extract_data_from_pdf(self.full_pdf_path)
            if not self.all_data.get("products"):
                messagebox.showerror("Error de Lectura", "No se pudieron extraer productos del PDF.")
                return
            self.update_summary_text()
            self._fill_editable_fields()
            # Insertar s25 entre s2 y s3 usando place en el canvas frame
            self.s25.pack(fill=tk.X, padx=10, pady=4,
                          after=self.summary_text.master)
            productos = sorted(set(p['producto'] for p in self.all_data['products']))
            for prod in productos:
                self.add_product_widget(prod)
            self.status_label['text']     = f"Se encontraron {len(productos)} producto(s). Asigne plantillas."
            self._set_step_done(2, True)
            embarque_v = self.edit_embarque.get().strip() or self.all_data["general_info"].get("embarque","")
            po_v       = self.edit_po.get().strip()       or self.all_data["general_info"].get("po","")
            self.work_label.config(
                text=f"  Trabajando en:  {embarque_v}  ·  PO {po_v}  " if embarque_v else "")
            if embarque_v:
                self._check_embarque_duplicado(embarque_v, po_v)
            self.view_pdf_button['state'] = tk.NORMAL
        except Exception as e:
            messagebox.showerror("Error Crítico", f"Error al procesar el PDF:\n{e}")

    def _fill_editable_fields(self):
        """Precarga los campos editables con los datos extraídos del PDF."""
        info = self.all_data.get("general_info", {})
        self.edit_embarque.set(info.get("embarque", ""))
        self.edit_cliente.set(info.get("cliente", ""))
        self.edit_po.set(info.get("po", ""))
        self.edit_contenedor.set(info.get("contenedor", ""))
        self.edit_sello.set(info.get("sello", ""))
        self.edit_recorder.set(info.get("recorder", ""))
        # Detectar palletized desde load_type extraído
        load_raw = info.get("load_type", "")
        tipo     = detect_palletized(load_raw) or "pallet"
        self.edit_palletized.set(tipo)

        # Construir campos Brix/pH + selector calidad + defectos por producto
        for w in self.brix_ph_frame.winfo_children():
            w.destroy()
        self.brix_ph_vars    = {}
        self.quality_formato = {}
        self.quality_vars    = {}
        productos = sorted(set(p["producto"] for p in self.all_data.get("products", [])))
        quality_names = list(config["quality_formats"].keys())

        cliente_actual = self.all_data.get("general_info", {}).get("cliente", "")
        es_camerican   = is_camerican(cliente_actual)

        for i, prod in enumerate(productos):
            # Separador por producto
            sep = ttk.LabelFrame(self.brix_ph_frame,
                                 text=prod[:50], padding=6)
            sep.pack(fill=tk.X, pady=4, padx=4)

            if es_camerican:
                # ── MODO CAMERICAN ──────────────────────────────────────────
                # Leer labels de Brix/pH desde el template
                # Intentar leer template si ya fue asignado en Paso 3
                template_path = self.product_widgets.get(prod, {}).get("path", "")
                cam_brix_labels = []
                cam_frutas      = []
                cam_params      = []
                if template_path and os.path.exists(template_path):
                    try:
                        tdoc = Document(template_path)
                        cam_brix_labels = read_camerican_brix_labels(tdoc)
                        cam_frutas, cam_params, _, cam_defaults, cam_limits = read_camerican_defect_structure(tdoc)
                    except Exception as e:
                        logging.warning(f"Camerican read error: {e}")
                        cam_defaults = {}
                        cam_limits   = {}

                # Campos Brix/pH dinámicos
                brix_ph_cam = {}
                row_bp = ttk.Frame(sep)
                row_bp.pack(fill=tk.X, pady=(0,4))
                for label_orig, label_key in cam_brix_labels:
                    v = tk.StringVar()
                    brix_ph_cam[label_orig] = v
                    ttk.Label(row_bp, text=label_orig + ":", anchor="e",
                              width=max(20, len(label_orig)+2)).pack(side=tk.LEFT, padx=(0,4))
                    ttk.Entry(row_bp, textvariable=v, width=8).pack(side=tk.LEFT, padx=(0,12))

                self.brix_ph_vars[prod] = {"camerican": brix_ph_cam}

                # Tabla de defectos Camerican
                if cam_frutas and cam_params:
                    ttk.Label(sep, text="Resultados de Defectos (average):",
                              font=("Helvetica", 8, "bold")).pack(anchor=tk.W, pady=(4,2))
                    cam_frame = ttk.Frame(sep)
                    cam_frame.pack(fill=tk.X)
                    ttk.Label(cam_frame, text="Parámetro", width=22,
                              font=("Helvetica", 8, "bold")).grid(row=0, column=0, padx=2, sticky="w")
                    for ci, fruta in enumerate(cam_frutas):
                        ttk.Label(cam_frame, text=fruta, width=12,
                                  font=("Helvetica", 8, "bold")).grid(row=0, column=ci+1, padx=2)
                    cam_matrix = {}
                    for ri, param in enumerate(cam_params):
                        ttk.Label(cam_frame, text=param, width=22,
                                  anchor="w").grid(row=ri+1, column=0, padx=2, pady=1, sticky="w")
                        cam_matrix[param] = {}
                        for ci, fruta in enumerate(cam_frutas):
                            default_val = cam_defaults.get(param, {}).get(fruta, "")
                            v = tk.StringVar(value=default_val)
                            cam_matrix[param][fruta] = v
                            entry = ttk.Entry(cam_frame, textvariable=v, width=10)
                            entry.grid(row=ri+1, column=ci+1, padx=2, pady=1)
                            lim_num, lim_tipo, lim_txt = cam_limits.get(param, {}).get(fruta, (None, None, ""))
                            def _make_v2(e, sv, ln, lt2):
                                def _chk(*_):
                                    vt = sv.get().strip()
                                    if ln is None or lt2.lower() in ("none","n/a","na",""):
                                        e.configure(foreground="black"); return
                                    m2 = re.search(r"(\d+\.?\d*)", vt)
                                    if not m2:
                                        e.configure(foreground="black"); return
                                    try:
                                        e.configure(foreground="red" if float(m2.group(1)) > ln else "black")
                                    except Exception:
                                        e.configure(foreground="black")
                                sv.trace_add("write", _chk); _chk()
                            _make_v2(entry, v, lim_num, lim_txt)
                    self.quality_vars[prod]    = {"__camerican_matrix__": cam_matrix,
                                                   "__frutas__": cam_frutas}
                    self.quality_formato[prod] = tk.StringVar(value="__camerican__")
                else:
                    self.quality_vars[prod]    = {}
                    self.quality_formato[prod] = tk.StringVar(value="")

            else:
                # ── MODO NORMAL ─────────────────────────────────────────────
                # Fila Brix + pH
                row_bp = ttk.Frame(sep)
                row_bp.pack(fill=tk.X, pady=(0,4))
                brix_var = tk.StringVar()
                ph_var   = tk.StringVar()
                ttk.Label(row_bp, text="Average Brix:", width=16, anchor="e").pack(side=tk.LEFT, padx=(0,4))
                ttk.Entry(row_bp, textvariable=brix_var, width=10).pack(side=tk.LEFT, padx=(0,16))
                ttk.Label(row_bp, text="Average pH:", width=14, anchor="e").pack(side=tk.LEFT, padx=(0,4))
                ttk.Entry(row_bp, textvariable=ph_var, width=10).pack(side=tk.LEFT)
                self.brix_ph_vars[prod] = {"brix": brix_var, "ph": ph_var}

                # Fila selector formato calidad
                row_fmt = ttk.Frame(sep)
                row_fmt.pack(fill=tk.X, pady=(0,4))
                auto_fmt = _detect_quality_format(prod)
                fmt_var = tk.StringVar(value=auto_fmt if auto_fmt else (quality_names[0] if quality_names else ""))
                self.quality_formato[prod] = fmt_var
                ttk.Label(row_fmt, text="Formato calidad:", width=16, anchor="e").pack(side=tk.LEFT, padx=(0,4))
                fmt_combo = ttk.Combobox(row_fmt, textvariable=fmt_var,
                                         values=quality_names, state="readonly", width=24)
                fmt_combo.pack(side=tk.LEFT, padx=(0,8))
                ttk.Button(row_fmt, text="Cargar campos",
                           command=lambda p=prod: self._build_quality_fields(p)).pack(side=tk.LEFT)

                # Frame para campos de defectos
                fields_frame = ttk.Frame(sep)
                fields_frame.pack(fill=tk.X)
                sep._quality_fields_frame = fields_frame
                self.quality_vars[prod] = {}

                # Cargar campos automáticamente con el primer formato
                self._build_quality_fields(prod)

    def update_summary_text(self):
        self.summary_text.config(state=tk.NORMAL)
        self.summary_text.delete(1.0, tk.END)
        info = self.all_data.get("general_info", {})
        self.summary_text.insert(tk.END,
            f"Cliente:    {info.get('cliente', 'No encontrado')}\n"
            f"Embarque:   {info.get('embarque', 'No encontrado')}\n"
            f"PO Número:  {info.get('po', 'No encontrado')}\n"
            f"Contenedor: {info.get('contenedor', 'No encontrado')}"
        )
        self.summary_text.config(state=tk.DISABLED)

    @staticmethod
    def _es_mix(producto):
        """Devuelve True si el nombre del producto indica que es un mix."""
        p = producto.lower()
        # Palabras clave directas de mix
        mix_words = ["mix", "blend", "medley", "mixed", "assorted",
                     "variety", "tropical", "berry blend", "fruit blend"]
        if any(w in p for w in mix_words):
            return True
        # Dos o más frutas separadas por conectores
        conectores = [" & ", " and ", " + ", " con ", " / ", " with "]
        if any(c in p for c in conectores):
            return True
        return False

    def _find_template_auto(self, producto):
        """Busca recursivamente un template que coincida con producto/cliente.
        Si el producto es un mix, busca Template_Mix en la carpeta del cliente."""
        template_folder = config.get("template_folder", "")
        if not template_folder or not os.path.isdir(template_folder):
            return ""
        cliente = self.all_data.get("general_info", {}).get("cliente", "").lower()
        prod_lower = producto.lower()

        # ── Caso Mix: buscar Template_Mix en carpeta del cliente ─────────
        if self._es_mix(producto):
            cliente_words = [w for w in re.split(r"[\s\-_/]+", cliente)
                             if len(w) > 2]
            for dirpath, dirs, files in os.walk(template_folder):
                folder_low = dirpath.lower()
                # La carpeta debe corresponder al cliente
                if not any(w in folder_low for w in cliente_words):
                    continue
                for fname in files:
                    if (fname.lower().startswith("template_mix")
                            and fname.lower().endswith(".docx")):
                        return os.path.join(dirpath, fname)

        # ── Caso normal: buscar por palabras clave del producto ──────────
        ignore = {"iqf", "frozen", "organic", "conventional", "the", "and", "with",
                  "de", "del", "para", "por", "fresh", "sliced", "blend", "mixed",
                  "mix", "berry", "fruit", "assorted"}
        prod_words = [w for w in re.split(r"[\s\-_/&+]+", prod_lower)
                      if len(w) > 2 and w not in ignore]
        cliente_words = [w for w in re.split(r"[\s\-_/]+", cliente)
                         if len(w) > 2 and w not in ignore]

        best_path  = ""
        best_score = 0
        for dirpath, dirs, files in os.walk(template_folder):
            for fname in files:
                if not fname.lower().endswith(".docx"):
                    continue
                fpath      = os.path.join(dirpath, fname)
                fname_low  = fname.lower()
                folder_low = dirpath.lower()
                combined   = fname_low + " " + folder_low
                score = 0
                for w in prod_words:
                    if w in combined:
                        score += 2
                for w in cliente_words:
                    if w in combined:
                        score += 1
                if score > best_score:
                    best_score = score
                    best_path  = fpath
        return best_path if best_score >= 2 else ""

    def add_product_widget(self, producto):
        pf = ttk.Frame(self.products_frame)
        pf.pack(fill=tk.X, pady=2)
        nombre = producto[:40] + ("..." if len(producto) > 40 else "")
        btn = ttk.Button(pf, text=f"Asignar Plantilla para '{nombre}'",
                         command=lambda p=producto: self.select_template(p))
        btn.pack(side=tk.LEFT)
        lv = tk.StringVar(value="Buscando plantilla...")
        lbl = ttk.Label(pf, textvariable=lv, foreground="#2A4A6B")
        lbl.pack(side=tk.LEFT, padx=10)
        self.product_widgets[producto] = {
            'button': btn, 'label_var': lv, 'path': '',
            'micro_formato': list(config["micro_formats"].keys())[0]
        }
        # Intentar auto-detectar template
        auto_path = self._find_template_auto(producto)
        if auto_path:
            self.product_widgets[producto]['path'] = auto_path
            lv.set(f"\u2714 Auto: {os.path.basename(auto_path)}  —  📂 {os.path.basename(os.path.dirname(auto_path))}")
            btn.configure(style="Success.TButton")
            cliente_actual = self.all_data.get("general_info", {}).get("cliente", "")
            if is_camerican(cliente_actual):
                self._refresh_camerican_fields(producto, auto_path)
            self.check_if_ready_to_generate()
        else:
            lv.set("Ninguna plantilla seleccionada.")

    def select_template(self, producto):
        path = filedialog.askopenfilename(
            title=f"Selecciona el template para: {producto}",
            filetypes=[("Word files", "*.docx")],
            initialdir=config.get("template_folder", "")
        )
        if path:
            self.product_widgets[producto]['path'] = path
            carpeta = os.path.basename(os.path.dirname(path))
            self.product_widgets[producto]['label_var'].set(
                f"{os.path.basename(path)}  \u2014  \U0001f4c2 {carpeta}")
            self.product_widgets[producto]['button'].configure(style="Success.TButton")
            self.check_if_ready_to_generate()
            # Si es Camerican y el paso 2.5 ya fue construido, refrescar campos
            cliente_actual = self.all_data.get("general_info", {}).get("cliente", "")
            if is_camerican(cliente_actual):
                self._refresh_camerican_fields(producto, path)

    def _refresh_camerican_fields(self, producto, template_path):
        """Reconstruye los campos Brix/defectos Camerican cuando se asigna el template."""
        try:
            tdoc            = Document(template_path)
            cam_brix_labels = read_camerican_brix_labels(tdoc)
            cam_frutas, cam_params, _, cam_defaults, cam_limits = read_camerican_defect_structure(tdoc)
        except Exception as e:
            logging.warning(f"_refresh_camerican_fields error: {e}")
            return

        for child in self.brix_ph_frame.winfo_children():
            if hasattr(child, 'cget') and child.cget('text').strip() == producto[:50].strip():
                for w in child.winfo_children():
                    w.destroy()

                # Campos Brix/pH
                brix_ph_cam = {}
                row_bp = ttk.Frame(child)
                row_bp.pack(fill=tk.X, pady=(0,4))
                for label_orig, _ in cam_brix_labels:
                    v = tk.StringVar()
                    brix_ph_cam[label_orig] = v
                    ttk.Label(row_bp, text=label_orig + ":", anchor="e",
                              width=max(20, len(label_orig)+2)).pack(side=tk.LEFT, padx=(0,4))
                    ttk.Entry(row_bp, textvariable=v, width=8).pack(side=tk.LEFT, padx=(0,12))
                self.brix_ph_vars[producto] = {"camerican": brix_ph_cam}

                # Tabla de defectos con defaults y validación fuera de spec
                if cam_frutas and cam_params:
                    ttk.Label(child, text="Resultados de Defectos (average):",
                              font=("Helvetica", 8, "bold")).pack(anchor=tk.W, pady=(4,2))
                    cam_frame = ttk.Frame(child)
                    cam_frame.pack(fill=tk.X)
                    ttk.Label(cam_frame, text="Parámetro", width=22,
                              font=("Helvetica", 8, "bold")).grid(row=0, column=0, padx=2, sticky="w")
                    for ci, fruta in enumerate(cam_frutas):
                        ttk.Label(cam_frame, text=fruta, width=12,
                                  font=("Helvetica", 8, "bold")).grid(row=0, column=ci+1, padx=2)
                    cam_matrix = {}
                    for ri, param in enumerate(cam_params):
                        ttk.Label(cam_frame, text=param, width=22,
                                  anchor="w").grid(row=ri+1, column=0, padx=2, pady=1, sticky="w")
                        cam_matrix[param] = {}
                        for ci, fruta in enumerate(cam_frutas):
                            default_val = cam_defaults.get(param, {}).get(fruta, "")
                            v = tk.StringVar(value=default_val)
                            cam_matrix[param][fruta] = v
                            entry = ttk.Entry(cam_frame, textvariable=v, width=10)
                            entry.grid(row=ri+1, column=ci+1, padx=2, pady=1)

                            # Validación fuera de spec — cambiar color si supera límite
                            lim_num, lim_tipo, lim_txt = cam_limits.get(param, {}).get(fruta, (None, None, ""))
                            def _make_validator(e, sv, ln, lt, lt2):
                                def _check(*_):
                                    val_txt = sv.get().strip()
                                    if ln is None or lt2.lower() in ("none", "n/a", "na", ""):
                                        e.configure(foreground="black")
                                        return
                                    import re as _re3
                                    m = re.search(r"(\d+\.?\d*)", val_txt)
                                    if not m:
                                        e.configure(foreground="black")
                                        return
                                    try:
                                        val_num = float(m.group(1))
                                        if val_num > ln:
                                            e.configure(foreground="red")
                                        else:
                                            e.configure(foreground="black")
                                    except Exception:
                                        e.configure(foreground="black")
                                sv.trace_add("write", _check)
                                _check()
                            _make_validator(entry, v, lim_num, lim_tipo, lim_txt)

                    self.quality_vars[producto]    = {"__camerican_matrix__": cam_matrix,
                                                       "__frutas__": cam_frutas}
                    self.quality_formato[producto] = tk.StringVar(value="__camerican__")
                break

    def clear_products(self):
        self.open_folder_button.pack_forget()
        self.open_file_button.pack_forget()
        for w in self.products_frame.winfo_children():
            w.destroy()
        self.product_widgets  = {}
        self._micro_results   = {}
        self._micro_formatos  = {}
        self._productos_lotes = {}
        self._productos_agg   = {}
        self.summary_text.config(state=tk.NORMAL)
        self.summary_text.delete(1.0, tk.END)
        self.summary_text.config(state=tk.DISABLED)
        self.generate_button['state'] = tk.DISABLED
        self.micro_button['state']    = tk.DISABLED
        self.micro_status_label['text'] = "Sin resultados ingresados."
        self.micro_status_label['foreground'] = "#2A4A6B"

    # ── SESION ─────────────────────────────────────────────────────
    def save_session(self):
        """Guarda el estado actual para restaurarlo la proxima vez."""
        try:
            session = {
                "pdf_path":       self.full_pdf_path,
                "output_folder":  self.output_folder.get(),
                "embarque":       self.edit_embarque.get(),
                "cliente":        self.edit_cliente.get(),
                "po":             self.edit_po.get(),
                "contenedor":     self.edit_contenedor.get(),
                "sello":          self.edit_sello.get(),
                "recorder":       self.edit_recorder.get(),
                "palletized":     self.edit_palletized.get(),
                "templates":      {p: w["path"] for p, w in self.product_widgets.items()},
                "micro_formatos": self._micro_formatos,
                "micro_results":  self._micro_results,
            }
            save_session_data(session, APP_DIR)
        except Exception:
            pass

    def load_session(self):
        """Carga sesion guardada. Devuelve dict o None."""
        try:
            return load_session_data(APP_DIR)
        except Exception:
            pass
        return None

    def restore_session(self, session):
        """Intenta restaurar una sesion guardada."""
        pdf = session.get("pdf_path", "")
        if not pdf or not os.path.exists(pdf):
            return
        emb = session.get("embarque", "")
        cli = session.get("cliente", "")
        po  = session.get("po", "")
        resp = messagebox.askyesno(
            "Restaurar sesion",
            "Sesion guardada encontrada:\n\n"
            f"  PDF:      {os.path.basename(pdf)}\n"
            f"  Embarque: {emb}\n"
            f"  Cliente:  {cli}\n\n"
            "Deseas continuar donde lo dejaste?"
        )
        if not resp:
            return
        self.full_pdf_path = pdf
        self.pdf_path.set(os.path.basename(pdf))
        self.process_button["state"]  = tk.NORMAL
        self.view_pdf_button["state"] = tk.NORMAL
        self.output_folder.set(session.get("output_folder", ""))
        try:
            self.all_data = extract_data_from_pdf(pdf)
        except Exception:
            messagebox.showerror("Error", "No se pudo recargar el PDF de la sesion.")
            return
        self.edit_embarque.set(emb)
        self.edit_cliente.set(cli)
        self.edit_po.set(po)
        self.edit_contenedor.set(session.get("contenedor", ""))
        self.edit_sello.set(session.get("sello", ""))
        self.edit_recorder.set(session.get("recorder", ""))
        self.edit_palletized.set(session.get("palletized", "pallet"))
        self.s25.pack(fill=tk.X, padx=14, pady=5)
        self._fill_editable_fields()
        self.update_summary_text()
        templates = session.get("templates", {})
        productos = sorted({p["producto"] for p in self.all_data.get("products", [])})
        for prod in productos:
            self.add_product_widget(prod)
            path = templates.get(prod, "")
            if path and os.path.exists(path):
                self.product_widgets[prod]["path"] = path
                self.product_widgets[prod]["label_var"].set("Plantilla: " + os.path.basename(path))
                self.product_widgets[prod]["button"].configure(style="Success.TButton")
        self._micro_formatos = session.get("micro_formatos", {})
        self._micro_results  = session.get("micro_results",  {})
        if self._micro_results:
            total = sum(1 for p in self._micro_results.values()
                        for l in p.values() for v in l.values() if v)
            self.micro_status_label["text"]       = f"Micro restaurada: {total} campo(s)."
            self.micro_status_label["foreground"] = "green"
        if emb:
            self.work_label.config(text=f"  Trabajando en:  {emb}  ·  PO {po}  ")
        self.micro_button["state"] = tk.NORMAL
        self.check_if_ready_to_generate()
        self.status_label["text"] = "  Sesion restaurada correctamente."

    def reset_all(self):
        """Limpia todo y vuelve al estado inicial."""
        if self.product_widgets or self.all_data:
            if not messagebox.askyesno("Nueva carga",
                    "¿Limpiar todo y empezar de nuevo?\n\nSe perderán los datos actuales."):
                return
        # Limpiar productos y estado
        self.clear_products()
        # Resetear PDF
        self.pdf_path.set("")
        self.full_pdf_path = ""
        self.all_data      = {}
        self.last_generated_file = None
        # Resetear campos editables
        for var in [self.edit_embarque, self.edit_cliente, self.edit_po,
                    self.edit_contenedor, self.edit_sello, self.edit_recorder]:
            var.set("")
        self.edit_palletized.set("pallet")
        # Limpiar brix/ph y calidad
        for w in self.brix_ph_frame.winfo_children():
            w.destroy()
        self.brix_ph_vars    = {}
        self.quality_formato = {}
        self.quality_vars    = {}
        # Ocultar paso 2.5
        self.s25.pack_forget()
        # Resetear botones
        self.process_button['state']  = tk.DISABLED
        self.view_pdf_button['state'] = tk.DISABLED
        # Resetear work_label y status
        self.work_label.config(text="")
        self.status_label['text'] = "  Listo para empezar."
        # Scroll al tope
        self._gen_canvas.yview_moveto(0)

    def _set_step_done(self, step, done):
        """Marca un paso como completo (verde) o pendiente (normal)."""
        frame = self._step_frames.get(step)
        if not frame:
            return
        labels = {
            1: "  1 · Packing List",
            2: "  2 · Cargar Datos",
            3: "  3 · Plantillas por Producto",
            4: "  4 · Carpeta de Destino",
            5: "  5 · Microbiología",
            6: "  6 · Generar Archivos",
        }
        base = labels.get(step, "")
        if done:
            frame.configure(text=base + "  \u2714")
            try:
                style = ttk.Style()
                style.configure(f"Done{step}.TLabelframe.Label",
                    background="#C8DFF0", foreground="#1E7C42",
                    font=("Segoe UI", 9, "bold"))
                frame.configure(style=f"Done{step}.TLabelframe")
                style.configure(f"Done{step}.TLabelframe",
                    background="#C8DFF0", bordercolor="#1E7C42")
            except Exception:
                pass
        else:
            frame.configure(text=base)
            try:
                frame.configure(style="TLabelframe")
            except Exception:
                pass

    def check_if_ready_to_generate(self):
        all_set = all(w['path'] for w in self.product_widgets.values())
        # Actualizar checkmarks
        self._set_step_done(3, bool(all_set and self.product_widgets))
        self._set_step_done(4, bool(self.output_folder.get()))
        self._set_step_done(5, bool(self._micro_results))
        if self.output_folder.get() and all_set and self.product_widgets:
            self.generate_button['state'] = tk.NORMAL
            self.micro_button['state']    = tk.NORMAL
            self.status_label['text']     = "  \u2714 Listo para generar."
        else:
            self.generate_button['state'] = tk.DISABLED

    def _prepare_lotes(self):
        """Agrupa lotes por producto. Llena _productos_lotes y _productos_agg."""
        self._productos_lotes = {}
        self._productos_agg   = {}
        for producto, widgets in self.product_widgets.items():
            if not widgets['path']: continue
            df = [p for p in self.all_data['products'] if p['producto'] == producto]
            if not df: continue
            agg = {}
            for lot in df:
                base = lot['lote'].split('-')[0]
                try:    qty = float(lot['cantidad'])
                except: qty = 0
                if base not in agg:
                    agg[base] = {'cantidad_total': qty, 'fecha_produccion': lot['fecha_produccion']}
                else:
                    agg[base]['cantidad_total'] += qty
            self._productos_lotes[producto] = list(agg.keys())
            self._productos_agg[producto]   = agg

    def open_micro_dialog(self):
        """Paso 5: abre la ventana de microbiología para ingresar o revisar."""
        self._prepare_lotes()
        cliente_actual = self.all_data["general_info"].get("cliente", "")
        # Si ya se eligió un formato antes, usarlo; si no, auto-detectar por cliente
        productos_formatos = {
            p: self._micro_formatos[p] if p in self._micro_formatos
               else _detect_micro_format(cliente_actual, p)
            for p in self.product_widgets
        }
        micro_history = load_micro_history()
        dialogo = MicrobiologyInputDialog(
            self.root, self._productos_lotes, productos_formatos,
            cliente_actual, micro_history
        )
        # Pre-llenar con resultados ya ingresados si los hay
        self.root.wait_window(dialogo)

        if dialogo.confirmed:
            self._micro_results  = dialogo.get_results()
            self._micro_formatos = dialogo.get_formatos()
            # Contar cuántos campos tienen valor
            total = sum(
                1 for p in self._micro_results.values()
                for l in p.values()
                for v in l.values() if v
            )
            self.micro_status_label['text']       = f"✔ {total} resultado(s) ingresado(s)."
            self.micro_status_label['foreground'] = "green"
            self._set_step_done(5, True)

    def _validar_antes_generar(self):
        """Chequea que todo esté listo. Devuelve lista de problemas."""
        problemas = []
        if not self.output_folder.get() or not os.path.isdir(self.output_folder.get()):
            problemas.append("\u26a0 Carpeta de destino no existe o no está configurada.")
        for prod, w in self.product_widgets.items():
            if not w['path']:
                problemas.append(f"\u26a0 '{prod[:40]}': sin plantilla asignada.")
            elif not os.path.exists(w['path']):
                problemas.append(f"\u26a0 '{prod[:40]}': plantilla no encontrada en disco.")
        if not self.product_widgets:
            problemas.append("\u26a0 No hay productos cargados.")
        return problemas

    def generate_coas(self):
        # Validar antes de empezar
        problemas = self._validar_antes_generar()
        if problemas:
            detalle = "\n".join(problemas)
            continuar = messagebox.askyesno(
                "Problemas detectados",
                f"Se encontraron los siguientes problemas:\n\n{detalle}\n\n"
                "¿Deseas continuar de todas formas?"
            )
            if not continuar:
                return
        self.status_label['text'] = "Generando archivos... por favor espere."
        self.open_folder_button.pack_forget()
        self.open_file_button.pack_forget()
        # Mostrar barra de progreso
        self.progress_bar.pack(fill=tk.X, side=tk.BOTTOM)
        self.progress_bar['value'] = 0
        self.root.update_idletasks()

        try:
            general_info   = self.all_data["general_info"]
            # Usar valores editados en Paso 2.5 (si el usuario corrigió algo)
            embarque       = self.edit_embarque.get().strip()  or general_info.get("embarque", "SIN_EMBARQUE")
            cliente_actual = self.edit_cliente.get().strip()   or general_info.get("cliente", "")
            # Actualizar general_info con los valores editados para que fill_all_tables los use
            general_info = dict(general_info)
            general_info["embarque"]   = embarque
            general_info["cliente"]    = cliente_actual
            general_info["po"]         = self.edit_po.get().strip()         or general_info.get("po", "")
            general_info["contenedor"] = self.edit_contenedor.get().strip() or general_info.get("contenedor", "")
            general_info["sello"]      = self.edit_sello.get().strip()      or general_info.get("sello", "")
            general_info["recorder"]   = self.edit_recorder.get().strip()   or general_info.get("recorder", "")
            general_info["palletized"] = self.edit_palletized.get()

            self._prepare_lotes()

            usar_micro = bool(self._micro_results)
            if not usar_micro:
                usar_micro = messagebox.askyesno(
                    "Microbiología",
                    "No has ingresado resultados microbiológicos.\n\n"
                    "¿Deseas ingresarlos ahora antes de generar?\n\n"
                    "• Sí → Abre la ventana de ingreso\n"
                    "• No → Genera sin llenar microbiología"
                )
                if usar_micro:
                    self.open_micro_dialog()
                    usar_micro = bool(self._micro_results)

            generated_files       = 0
            generated_files_paths = []
            registro_filas        = []
            fecha_gen             = datetime.now().strftime("%d/%m/%Y %H:%M")

            for producto, widgets in self.product_widgets.items():
                template_path = widgets['path']
                if not template_path or not os.path.exists(template_path):
                    messagebox.showwarning("Plantilla Inválida", f"No se encontró la plantilla para '{producto}'.")
                    continue

                df = [p for p in self.all_data['products'] if p['producto'] == producto]
                if not df: continue

                agg              = self._productos_agg.get(producto, {})
                final_lotes      = list(agg.keys())
                final_cantidades = [str(d['cantidad_total']) for d in agg.values()]
                final_fechas     = [d['fecha_produccion']    for d in agg.values()]

                first_row  = df[0]
                suma_cajas = sum(float(p['cantidad'])         for p in df)
                suma_neto  = sum(float(p['peso_neto_total'])  for p in df)
                suma_bruto = sum(float(p['peso_bruto_total']) for p in df)

                replacements = {
                    "Product Name":                  producto,
                    "Customer":                      cliente_actual,
                    "Shipment Number":               f"{general_info.get('embarque','')} / {general_info.get('po','')}",
                    "FCL Number":                    general_info.get("contenedor", ""),
                    "Seal Number":                   general_info.get("sello", ""),
                    "Ryan Recorder Number":          general_info.get("recorder", ""),
                    "Net Weight per unit":           str(first_row["peso_x_unidad"]),
                    "Net Weight per Container":      f"{suma_neto:.2f}",
                    "Gross Weight per Container":    f"{suma_bruto:.2f}",
                    "Number of units per Container": f"{int(suma_cajas)}",
                    "Manufacture Date":              ", ".join(final_fechas),
                    "Expiration Date":               "",
                }

                if is_special_customer(cliente_actual):
                    exp = []
                    for ds in final_fechas:
                        try:
                            pd_ = datetime.strptime(ds, '%d/%m/%Y')
                            exp.append(pd_.replace(year=pd_.year + 3).strftime('%d/%m/%Y'))
                        except ValueError:
                            exp.append("N/A")
                    replacements["Expiration Date"] = ", ".join(exp)

                # Agregar Brix y pH
                brix_ph     = self.brix_ph_vars.get(producto, {})
                _es_cam_gen = is_camerican(cliente_actual)

                if _es_cam_gen and "camerican" in brix_ph:
                    # Camerican: múltiples Brix/pH → llenar por label
                    cam_brix_vals = {k: v.get().strip()
                                     for k, v in brix_ph["camerican"].items()
                                     if v.get().strip()}
                else:
                    brix_val = brix_ph.get("brix", tk.StringVar()).get().strip()
                    ph_val   = brix_ph.get("ph",   tk.StringVar()).get().strip()
                    if brix_val:
                        replacements["Average Brix"] = brix_val
                    if ph_val:
                        replacements["Average pH"] = ph_val
                    cam_brix_vals = {}

                doc = Document(template_path)
                fill_all_tables(doc, replacements)

                # Camerican: llenar Brix/pH por label específico
                if cam_brix_vals:
                    fill_camerican_brix_ph(doc, cam_brix_vals)

                fill_lot_table(doc, {
                    "Lotes":       final_lotes,
                    "Cantidades":  final_cantidades,
                    "Fechas":      final_fechas,
                    "Cliente":     cliente_actual,
                    "MicroFormato": self._micro_formatos.get(producto, ""),
                })
                fill_palletized(doc, general_info.get("palletized", ""))
                replace_pesticide_text(doc, producto)

                # Calidad — defectos
                fmt_calidad  = self.quality_formato.get(producto, tk.StringVar()).get()
                q_vars       = self.quality_vars.get(producto, {})

                if fmt_calidad == "__camerican__" and "__camerican_matrix__" in q_vars:
                    # Camerican: llenar tabla de defectos con matrix
                    raw_matrix = q_vars["__camerican_matrix__"]
                    frutas_cam = q_vars.get("__frutas__", [])
                    quality_matrix = {param: {fruta: v.get().strip()
                                              for fruta, v in f_dict.items()}
                                      for param, f_dict in raw_matrix.items()}
                    fill_camerican_defects(doc, frutas_cam, quality_matrix)
                else:
                    quality_values = {k: v.get() for k, v in q_vars.items()
                                      if not k.startswith("__extra_")}
                    if fmt_calidad and quality_values:
                        fill_defects_table(doc, fmt_calidad, quality_values)

                # MICROBIOLOGÍA
                if usar_micro:
                    micro_per_lote = self._micro_results.get(producto, {})
                    fmt_nombre     = self._micro_formatos.get(producto, widgets['micro_formato'])
                    fmt            = config["micro_formats"].get(fmt_nombre, {})
                    tipo_micro     = fmt.get("tipo", "simple")
                    micro_first    = micro_per_lote.get(final_lotes[0], {}) if final_lotes else {}

                    if tipo_micro == "camerican":
                        duplicate_camerican_tables(doc, final_lotes, micro_per_lote)
                    elif tipo_micro == "n_series":
                        duplicate_n_series_tables(doc, final_lotes, micro_per_lote)
                        # Tablita HepA/Norovirus — siempre en primera página, nunca duplicar
                        micro_tables = find_micro_tables(doc)
                        for _, table, tipo_tabla in micro_tables:
                            if tipo_tabla == MICRO_TYPE_SMALL:
                                fill_micro_small(table, final_lotes, micro_per_lote)
                    else:
                        micro_tables = find_micro_tables(doc)
                        for _, table, tipo_tabla in micro_tables:
                            if tipo_tabla == MICRO_TYPE_SIMPLE:
                                fill_micro_simple(doc, table, final_lotes, micro_per_lote)
                            elif tipo_tabla == MICRO_TYPE_SMALL:
                                # SOLO llenar, NUNCA duplicar — columnas por lote
                                fill_micro_small(table, final_lotes, micro_per_lote)

                    # Guardar en historial
                    for lote in final_lotes:
                        lote_data = micro_per_lote.get(lote, {})
                        if any(lote_data.values()):
                            save_micro_history_record(cliente_actual, producto, lote, lote_data, fmt_nombre)

                filename    = estilizar_nombre_archivo(embarque, general_info.get('po', 'S_PO'), producto)
                output_path = os.path.join(self.output_folder.get(), filename)
                doc.save(output_path)
                generated_files += 1
                generated_files_paths.append(output_path)
                # Actualizar barra de progreso
                pct = int(generated_files / max(len(self.product_widgets), 1) * 100)
                self.progress_bar['value'] = pct
                self.root.update_idletasks()

                registro_filas.append([
                    fecha_gen, embarque, cliente_actual,
                    general_info.get('po', ''), producto,
                    ", ".join(final_lotes), int(suma_cajas), filename
                ])

            if generated_files > 0:
                registrar_coa(registro_filas)
                self.progress_bar['value'] = 100
                self.root.update_idletasks()
                self.root.after(1200, lambda: self.progress_bar.pack_forget())
                self.status_label['text'] = f"¡Proceso completado! Se generaron {generated_files} COA(s)."
                self.open_folder_button.pack(side=tk.RIGHT, padx=5)
                if generated_files == 1:
                    self.last_generated_file = generated_files_paths[0]
                    self.open_file_button.pack(side=tk.RIGHT, padx=5)
                messagebox.showinfo("Proceso Completado",
                    f"Se generaron {generated_files} COA(s) para el embarque {embarque}.\n"
                    f"Registro actualizado en '{COA_REGISTRY}'.")
            else:
                self.status_label['text'] = "Proceso finalizado. No se generó ningún archivo."

        except Exception as e:
            error_msg = traceback.format_exc()
            logging.error("Excepción no controlada:\n" + error_msg)
            self.progress_bar.pack_forget()
            self.status_label['text'] = "Error durante la generación. Revisa error_log.txt"
            messagebox.showerror("Error Crítico",
                f"Ocurrió un error inesperado:\n\n{e}\n\nSe guardó el detalle en 'error_log.txt'.")

# ============================================================
# PUNTO DE ENTRADA
# ============================================================
if __name__ == "__main__":
    root = tk.Tk()
    root.title("Generador de COA - Control de Calidad")
    
    # ESTA ES LA LÍNEA CLAVE:
    try:
        root.iconbitmap("ico.ico") 
    except:
        pass # Si no encuentra el archivo, usa el de defecto para no dar error
        
    app = COAGeneratorApp(root)
    root.protocol("WM_DELETE_WINDOW", lambda: (app.save_session(), root.destroy()))
    session = app.load_session()
    if session:
        root.after(400, lambda: app.restore_session(session))
    root.mainloop()
